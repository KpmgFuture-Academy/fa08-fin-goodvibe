"""HWPX heading / appendix pre-parse helpers for RAG admin flows."""
from __future__ import annotations

import html
import io
import re
import tempfile
import unicodedata
import zipfile
from dataclasses import dataclass
from pathlib import Path
from typing import Any

import lxml.etree as ET

MAX_MATCHED_SAMPLE_ITEMS = 20
_SYMBOL_CANDIDATE_RE = re.compile(r"^\s*([①-⑳㉮-㉻❍○●⊙■□▣◇◆▶▷☞·•◦◉◈▪▫◾◽◼◻◯❑])(?:\s+|$)")
_GENERIC_HEADING_PREFIX_RE = re.compile(
    r"^\s*(?:"
    r"[ⅠⅡⅢⅣⅤⅥⅦⅧⅨⅩⅪⅫIVXLCDM]+(?:\s+|$)"
    r"|[0-9]+[.)]?(?:\s+|$)"
    r"|[가-힣][.)](?:\s+|$)"
    r"|\([가-힣]\)(?:\s+|$)"
    r"|[①-⑳㉮-㉻❍○●⊙■□▣◇◆▶▷☞※·•◦◉◈▪▫◾◽◼◻◯❑](?:\s+|$)"
    r")"
)


@dataclass
class HwpxHeadingParseResult:
    """HWPX pre-parse 결과를 화면/서비스 계층에 전달하는 DTO."""

    source_lines: list[dict[str, Any]]
    main_source_lines: list[dict[str, Any]]
    appendix_source_lines: list[dict[str, Any]]
    heading_rows: list[dict[str, Any]]
    appendix_heading_rows: list[dict[str, Any]]


@dataclass
class ParsedBlock:
    """문서 ingest 단계의 공통 텍스트 블록 표현."""

    title: str
    text: str
    section: str = ""


class HwpxParsingError(Exception):
    """문서 pre-parse / ingest 에서 사용되는 공통 예외."""


class HwpxHeadingParser:
    """Parse HWPX documents and build heading / appendix candidates."""

    def __init__(
        self,
        *,
        main_schema: dict[str, Any],
        appendix_schema: dict[str, Any] | None = None,
        body_exit_criteria: dict[str, Any] | None = None,
        appendix_exit_criteria: dict[str, Any] | None = None,
    ) -> None:
        """본문/부록 schema 와 전환 기준을 바탕으로 parser 상태를 초기화한다."""
        self.main_schema = main_schema
        self.appendix_schema = appendix_schema
        self.body_exit_criteria = body_exit_criteria
        self.appendix_exit_criteria = appendix_exit_criteria

    def parse_content(self, content: bytes) -> HwpxHeadingParseResult:
        """HWPX 바이트를 읽어 pre-parse 라인과 heading 후보 요약을 생성한다."""
        source_lines = self._extract_hwpx_preparse_lines(content)
        main_source_lines, appendix_source_lines = self._split_preparse_lines_by_domain(source_lines)
        return HwpxHeadingParseResult(
            source_lines=source_lines,
            main_source_lines=main_source_lines,
            appendix_source_lines=appendix_source_lines,
            heading_rows=self._build_heading_rows(self.main_schema, raw_lines=main_source_lines),
            appendix_heading_rows=(
                self._build_heading_rows(
                    self.appendix_schema,
                    raw_lines=appendix_source_lines,
                    allow_new_candidates=False,
                )
                if isinstance(self.appendix_schema, dict)
                else []
            ),
        )

    def _extract_hwpx_preparse_lines(self, content: bytes) -> list[dict[str, Any]]:
        """HWPX XML 블록을 순회하며 paragraph/table 중심 pre-parse 라인을 만든다."""
        lines: list[dict[str, Any]] = []
        seen: set[tuple[str, str]] = set()
        seen_heading_keys: set[str] = set()

        def append_once(
            value: str,
            *,
            location: str,
            row_count: int | None = None,
            cell_count: int | None = None,
            metadata: dict[str, Any] | None = None,
        ) -> None:
            """중복을 제거하면서 pre-parse line 1건을 누적한다."""
            normalized = self._normalize_spaces(value)
            normalized_location = self._normalize_source_location(location)
            key = (normalized, normalized_location)
            heading_key = self._normalize_sample_key(normalized)
            if (
                normalized
                and heading_key
                and self._looks_like_generic_heading_text(normalized)
                and heading_key in seen_heading_keys
            ):
                return
            if not normalized or key in seen:
                return
            seen.add(key)
            if heading_key and self._looks_like_generic_heading_text(normalized):
                seen_heading_keys.add(heading_key)
            item: dict[str, Any] = {"text": normalized, "location": normalized_location}
            if row_count is not None:
                item["row_count"] = row_count
            if cell_count is not None:
                item["cell_count"] = cell_count
            if row_count == 1 and cell_count == 1:
                item["is_box"] = True
            if isinstance(metadata, dict):
                item.update(metadata)
            lines.append(item)

        parser = ET.XMLParser(resolve_entities=False, no_network=True, recover=True, huge_tree=True)

        with zipfile.ZipFile(io.BytesIO(content)) as archive:
            section_files = sorted(
                [name for name in archive.namelist() if re.match(r"^Contents/section\d+\.xml$", name)],
                key=self._sort_section_key,
            )
            for section_index, section_name in enumerate(section_files):
                xml_bytes = archive.read(section_name)
                try:
                    root = ET.fromstring(xml_bytes, parser=parser)
                except ET.XMLSyntaxError:
                    # XML 복구가 어려운 경우에도 paragraph 텍스트는 최대한 살려 preview 와 후보 탐색에 활용한다.
                    xml_text = xml_bytes.decode("utf-8", errors="ignore")
                    xml_text_without_tables = re.sub(
                        r"<[^>]+:tbl\b.*?</[^>]+:tbl>",
                        "",
                        xml_text,
                        flags=re.DOTALL,
                    )
                    paragraph_blocks = re.findall(
                        r"<[^>]+:p\b.*?</[^>]+:p>",
                        xml_text_without_tables,
                        flags=re.DOTALL,
                    )
                    for paragraph_xml in paragraph_blocks:
                        text = self._extract_paragraph_text_from_xml(paragraph_xml)
                        if text:
                            append_once(
                                text,
                                location="paragraph",
                                metadata={
                                    "section_name": section_name,
                                    "section_index": section_index,
                                    "block_type": "paragraph",
                                    "parser_mode": "regex_fallback",
                                },
                            )
                    continue

                block_index = 0
                for node in root.iter():
                    local_name = self._localname(node)
                    if local_name == "p":
                        # 표 내부 문단은 table 블록 처리에서 따로 수집한다.
                        if self._is_inside_table(node):
                            continue
                        # HWPX 는 table 을 감싸는 wrapper paragraph 를 두는 경우가 있어
                        # 이 문단까지 살리면 table 제목이 paragraph/table 로 중복 수집될 수 있다.
                        if self._is_table_wrapper_paragraph(node):
                            continue
                        text = self._normalize_spaces(self._extract_paragraph_text(node))
                        if text:
                            append_once(
                                text,
                                location="paragraph",
                                metadata={
                                    "section_name": section_name,
                                    "section_index": section_index,
                                    "block_index": block_index,
                                    "block_type": "paragraph",
                                    "xml_tag": local_name,
                                    "parser_mode": "xml",
                                    "para_shape_id": self._get_attr(node, "paraPrIDRef"),
                                    "style_id": self._get_attr(node, "styleIDRef"),
                                    "text_align": self._get_attr(node, "textDirection"),
                                },
                            )
                            block_index += 1
                        continue

                    if local_name not in {"tbl", "table"}:
                        continue

                    # 중첩 표는 상위 표 처리에서 이미 읽히므로 top-level table 만 대상으로 삼는다.
                    if self._is_inside_table(node):
                        continue

                    matrix = self._extract_table(node)
                    if not matrix:
                        block_index += 1
                        continue

                    row_count = len(matrix)
                    cell_count = max((len(row) for row in matrix), default=0)
                    flattened = self._flatten_preparse_table_row(matrix[0]) if row_count == 1 else ""
                    table_text = flattened or self._flatten_preparse_table_row(
                        [cell for row in matrix for cell in row if self._normalize_spaces(cell)]
                    )
                    if table_text:
                        append_once(
                            table_text,
                            location="table",
                            row_count=row_count,
                            cell_count=cell_count,
                            metadata={
                                "section_name": section_name,
                                "section_index": section_index,
                                "block_index": block_index,
                                "block_type": "table",
                                "xml_tag": local_name,
                                "parser_mode": "xml",
                                "table_rows": row_count,
                                "table_cells": sum(1 for row in matrix for cell in row if self._normalize_spaces(cell)),
                            },
                        )
                    block_index += 1
        return lines

    def _split_preparse_lines_by_domain(
        self,
        source_lines: list[dict[str, Any]],
    ) -> tuple[list[dict[str, Any]], list[dict[str, Any]]]:
        """본문/부록 schema 기준으로 pre-parse 라인을 domain 별로 분리한다."""
        if not isinstance(self.appendix_schema, dict):
            passthrough_lines: list[dict[str, Any]] = []
            current_main_depth: int | None = None
            for item in source_lines:
                if self._should_skip_main_preparse_item(item):
                    continue
                next_item = dict(item)
                matched_main_level = self._find_first_matching_level(item, self.main_schema)
                if matched_main_level:
                    current_main_depth = self._level_depth(matched_main_level)
                if current_main_depth is not None:
                    next_item["context_depth"] = current_main_depth
                next_item["domain_mode"] = "main"
                passthrough_lines.append(next_item)
            return passthrough_lines, []

        main_lines: list[dict[str, Any]] = []
        appendix_lines: list[dict[str, Any]] = []
        mode = "main"
        allow_main_reentry = True
        current_main_depth: int | None = None
        current_appendix_depth: int | None = None

        for item in source_lines:
            matched_main_level = self._find_first_matching_level(item, self.main_schema)
            matched_appendix_level = self._find_first_matching_level(item, self.appendix_schema)

            if mode == "main":
                if self._matches_exit_criteria(item, criteria=self.body_exit_criteria, target_schema=self.appendix_schema):
                    next_item = dict(item)
                    target_level = matched_appendix_level or self._find_first_matching_level(item, self.appendix_schema, depth=1)
                    current_appendix_depth = self._level_depth(target_level, 1) if target_level else 1
                    next_item["context_depth"] = current_appendix_depth
                    next_item["domain_mode"] = "appendix"
                    appendix_lines.append(next_item)
                    mode = "appendix"
                    allow_main_reentry = False
                    continue

                if matched_appendix_level and self._level_depth(matched_appendix_level) == 1:
                    next_item = dict(item)
                    current_appendix_depth = 1
                    next_item["context_depth"] = current_appendix_depth
                    next_item["domain_mode"] = "appendix"
                    appendix_lines.append(next_item)
                    mode = "appendix"
                    continue

                if self._should_skip_main_preparse_item(item):
                    continue

                next_item = dict(item)
                if matched_main_level:
                    current_main_depth = self._level_depth(matched_main_level)
                if current_main_depth is not None:
                    next_item["context_depth"] = current_main_depth
                next_item["domain_mode"] = "main"
                main_lines.append(next_item)
                continue

            if matched_appendix_level:
                next_item = dict(item)
                current_appendix_depth = self._level_depth(matched_appendix_level)
                next_item["context_depth"] = current_appendix_depth
                next_item["domain_mode"] = "appendix"
                appendix_lines.append(next_item)
                continue

            source_location = self._normalize_source_location(item.get("location") or "paragraph")

            matched_appendix_exit = (
                self._matches_exit_criteria(
                    item,
                    criteria=self.appendix_exit_criteria,
                    target_schema=self.main_schema,
                    restrict_box_main_reentry=True,
                    exclude_symbolic_reentry_levels=True,
                )
                if isinstance(self.appendix_exit_criteria, dict)
                else self._matches_default_appendix_exit(item, self.main_schema)
            )

            if matched_appendix_exit:
                next_item = dict(item)
                if matched_main_level:
                    current_main_depth = self._level_depth(matched_main_level)
                if current_main_depth is not None:
                    next_item["context_depth"] = current_main_depth
                next_item["domain_mode"] = "main"
                main_lines.append(next_item)
                mode = "main"
                continue

            if self._is_single_row_table_item(item) and matched_main_level and not self._is_symbolic_reentry_level(matched_main_level):
                next_item = dict(item)
                current_main_depth = self._level_depth(matched_main_level)
                next_item["context_depth"] = current_main_depth
                next_item["domain_mode"] = "main"
                main_lines.append(next_item)
                mode = "main"
                continue

            if source_location == "table" and not bool(item.get("is_box")):
                next_item = dict(item)
                if current_appendix_depth is not None:
                    next_item["context_depth"] = current_appendix_depth
                next_item["domain_mode"] = "appendix"
                appendix_lines.append(next_item)
                continue

            if not allow_main_reentry:
                next_item = dict(item)
                if current_appendix_depth is not None:
                    next_item["context_depth"] = current_appendix_depth
                next_item["domain_mode"] = "appendix"
                appendix_lines.append(next_item)
                continue

            next_item = dict(item)
            if current_appendix_depth is not None:
                next_item["context_depth"] = current_appendix_depth
            next_item["domain_mode"] = "appendix"
            appendix_lines.append(next_item)

        return main_lines, appendix_lines

    def _build_heading_rows(
        self,
        schema: dict[str, Any] | None,
        *,
        raw_lines: list[dict[str, Any]] | None = None,
        allow_new_candidates: bool = True,
    ) -> list[dict[str, Any]]:
        """schema 와 raw line 매칭 결과를 관리자 화면용 heading row 로 변환한다."""
        levels = schema.get("levels") if isinstance(schema, dict) else None
        if not isinstance(levels, list):
            return []

        rows: list[dict[str, Any]] = []
        existing_matchers: list[tuple[dict[str, Any], str, re.Pattern[str]]] = []
        occurrence_counts: dict[str, int] = {}
        matched_samples: dict[str, list[str]] = {}
        for idx, level in enumerate(levels, start=1):
            if not isinstance(level, dict):
                continue
            pattern = self._compile_level_pattern(level)
            notation = str(level.get("notation") or "").strip()
            notation_display = str(level.get("notation_display") or notation).strip() or notation
            row_id = f"existing-{idx}"
            occurrence_counts[row_id] = 0
            matched_samples[row_id] = []
            rule_type = str(level.get("rule_type") or "").strip()
            if rule_type == "appendix_title_table":
                existing_matchers.append((level, row_id, re.compile(r"^$")))
            elif pattern:
                try:
                    existing_matchers.append((level, row_id, re.compile(pattern)))
                except re.error:
                    pass
            rows.append(
                {
                    "row_id": row_id,
                    "depth": int(level.get("depth") or idx),
                    "rule_id": str(level.get("rule_id") or "").strip() or None,
                    "notation": notation,
                    "display_notation": notation_display,
                    "occurrence_count": 0,
                    "matched_samples": [],
                    "action": "삭제",
                    "pattern": pattern or None,
                    "rule_type": rule_type or None,
                    "rule_options": self._get_rule_options(level) or None,
                    "location": self._normalize_level_location(level.get("location") or "paragraph"),
                    "name": str(level.get("name") or "").strip() or None,
                    "is_new": False,
                }
            )

        candidate_counts: dict[str, int] = {}
        candidate_samples: dict[str, list[str]] = {}
        candidate_depths: dict[str, int] = {}
        candidate_first_seen: dict[str, int] = {}
        for line_index, item in enumerate(raw_lines or []):
            text = str(item.get("text") or "").strip()
            source_location = self._normalize_source_location(item.get("location") or "paragraph")
            is_box_table = bool(item.get("is_box")) and source_location == "table"
            matched_existing = False
            for level, row_id, compiled in existing_matchers:
                if is_box_table and self._normalize_level_location(level.get("location") or "paragraph") == "paragraph":
                    continue
                if not self._location_matches(str(level.get("location") or "paragraph"), source_location):
                    continue
                rule_type = str(level.get("rule_type") or "").strip()
                if rule_type == "appendix_title_table":
                    if self._matches_appendix_title_table(level, item):
                        if self._append_unique_sample(matched_samples[row_id], text):
                            occurrence_counts[row_id] = occurrence_counts.get(row_id, 0) + 1
                        matched_existing = True
                        break
                    continue
                if compiled.match(text) and self._is_heading_like_line(level, text):
                    if self._append_unique_sample(matched_samples[row_id], text):
                        occurrence_counts[row_id] = occurrence_counts.get(row_id, 0) + 1
                    matched_existing = True
                    break
            if matched_existing:
                continue

            fallback = self._fallback_match_existing(text, source_location, is_box_table=is_box_table, existing_matchers=existing_matchers)
            if fallback:
                _matched_level, row_id, _compiled = fallback
                if self._append_unique_sample(matched_samples[row_id], text):
                    occurrence_counts[row_id] = occurrence_counts.get(row_id, 0) + 1
                continue

            fallback_symbol_only = self._fallback_match_existing_symbol_only(
                text,
                source_location,
                is_box_table=is_box_table,
                existing_matchers=existing_matchers,
            )
            if fallback_symbol_only:
                _matched_level, row_id, _compiled = fallback_symbol_only
                if self._append_unique_sample(matched_samples[row_id], text):
                    occurrence_counts[row_id] = occurrence_counts.get(row_id, 0) + 1
                continue

            # 신규 후보는 paragraph 기반 symbol heading 만 보수적으로 추가한다.
            if not allow_new_candidates or source_location != "paragraph":
                continue
            if not self._is_heading_like_symbol_candidate(text):
                continue

            match = _SYMBOL_CANDIDATE_RE.match(text)
            if not match:
                continue
            symbol = match.group(1)
            context_depth = item.get("context_depth")
            inferred_depth = (int(context_depth) + 1) if isinstance(context_depth, int) and context_depth > 0 else self._guess_symbol_depth(symbol)
            candidate_depths[symbol] = max(candidate_depths.get(symbol, 0), inferred_depth)
            candidate_first_seen.setdefault(symbol, line_index)
            candidate_samples.setdefault(symbol, [])
            if self._append_unique_sample(candidate_samples[symbol], text):
                candidate_counts[symbol] = candidate_counts.get(symbol, 0) + 1

        for row in rows:
            row_id = str(row.get("row_id") or "")
            occurrence_count = occurrence_counts.get(row_id, 0)
            row["occurrence_count"] = occurrence_count
            row["matched_samples"] = matched_samples.get(row_id, [])
            row["action"] = "유지" if occurrence_count > 0 else "삭제"

        for idx, (symbol, count) in enumerate(
            sorted(candidate_counts.items(), key=lambda item: (candidate_depths.get(item[0], 0), candidate_first_seen.get(item[0], 999999), item[0])),
            start=1,
        ):
            rows.append(
                {
                    "row_id": f"candidate-{idx}",
                    "depth": candidate_depths.get(symbol) or self._guess_symbol_depth(symbol),
                    "rule_id": None,
                    "notation": symbol,
                    "display_notation": symbol,
                    "occurrence_count": count,
                    "matched_samples": candidate_samples.get(symbol, []),
                    "action": "신규",
                    "pattern": rf"^\s*{re.escape(symbol)}(?:\s+|$)",
                    "rule_type": "symbol",
                    "rule_options": {
                        "symbols": [symbol],
                        "allow_leading_space": True,
                        "require_space_or_eol": True,
                    },
                    "location": "paragraph",
                    "name": None,
                    "is_new": True,
                }
            )
        return rows

    @staticmethod
    def _guess_symbol_depth(symbol: str) -> int:
        """신규 symbol 후보의 기본 depth 를 추정한다."""
        if symbol in {"❍", "○", "●", "⊙", "■", "□", "▣", "◇", "◆", "❑", ""}:
            return 2
        if re.match(r"[①-⑳]", symbol):
            return 5
        if re.match(r"[㉮-㉻]", symbol):
            return 6
        return 2

    @staticmethod
    def _localname(node_or_tag: Any) -> str:
        """namespace 와 무관하게 element/tag 의 local name 을 반환한다."""
        if hasattr(node_or_tag, "tag"):
            return ET.QName(node_or_tag).localname
        tag = str(node_or_tag or "")
        return tag.split("}", 1)[-1] if "}" in tag else tag

    @staticmethod
    def _sort_section_key(path: str) -> int:
        """section 파일명을 숫자 기준으로 정렬하기 위한 key 를 만든다."""
        match = re.search(r"section(\d+)\.xml$", path)
        return int(match.group(1)) if match else 999999

    @staticmethod
    def _normalize_spaces(text: str) -> str:
        """연속 공백과 줄 구분을 화면/비교용 단일 공백 형태로 정규화한다."""
        return " ".join((text or "").split()).strip()

    @staticmethod
    def _strip_invisible_chars(text: str) -> str:
        """비교를 방해하는 invisible control 문자를 제거한다."""
        return "".join(ch for ch in (text or "") if unicodedata.category(ch) not in {"Cf", "Cc", "Cs"})

    @classmethod
    def _normalize_sample_key(cls, text: str) -> str:
        """heading 샘플 비교용 표준 key 를 만든다."""
        normalized = unicodedata.normalize("NFKC", cls._strip_invisible_chars(text))
        normalized = re.sub(r"\s*\|\s*", " ", normalized)
        normalized = re.sub(r"^([IVXLCDM]+)(?=[A-Za-z가-힣])", r"\1 ", normalized)
        normalized = re.sub(r"^([ⅠⅡⅢⅣⅤⅥⅦⅧⅨⅩⅪⅫ]+)(?=[A-Za-z가-힣])", r"\1 ", normalized)
        normalized = re.sub(r"^([0-9]+[.)]?)(?=[A-Za-z가-힣])", r"\1 ", normalized)
        return cls._normalize_spaces(normalized)

    @classmethod
    def _normalize_heading_display_text(cls, text: str) -> str:
        """사용자에게 보여줄 heading 샘플 텍스트를 원문 표기를 최대한 보존하며 정규화한다."""
        normalized = cls._strip_invisible_chars(text)
        normalized = re.sub(r"\s*\|\s*", " ", normalized)
        normalized = re.sub(r"^([IVXLCDM]+)(?=[A-Za-z가-힣])", r"\1 ", normalized)
        normalized = re.sub(r"^([ⅠⅡⅢⅣⅤⅥⅦⅧⅨⅩⅪⅫ]+)(?=[A-Za-z가-힣])", r"\1 ", normalized)
        normalized = re.sub(r"^([0-9]+[.)]?)(?=[A-Za-z가-힣])", r"\1 ", normalized)
        normalized = re.sub(r"^([가-힣][.)])(?=[A-Za-z가-힣])", r"\1 ", normalized)
        return cls._normalize_spaces(normalized)

    @classmethod
    def _normalize_sample_compare_key(cls, text: str) -> str:
        """중복 제거를 위한 비교 key 를 만든다."""
        return re.sub(r"\s+", "", cls._normalize_sample_key(text))

    @classmethod
    def _looks_like_generic_heading_text(cls, text: str) -> bool:
        """텍스트가 일반적인 heading marker 패턴으로 시작하는지 판정한다."""
        normalized = cls._normalize_sample_key(text)
        if not normalized:
            return False
        return bool(_GENERIC_HEADING_PREFIX_RE.match(normalized))

    @classmethod
    def _flatten_preparse_table_row(cls, row: list[str]) -> str:
        """표 1행을 preview/매칭용 단일 문자열로 평탄화한다."""
        return cls._normalize_spaces(" ".join(cell for cell in row if cls._normalize_spaces(cell)))

    @classmethod
    def _normalize_single_row_table_heading_text(cls, item: dict[str, Any], text: str) -> str:
        """1행 표 제목 후보를 비교 가능한 문자열로 정규화한다."""
        source_location = cls._normalize_source_location(item.get("location") or "paragraph")
        row_count = item.get("row_count")
        if source_location == "table" and row_count == 1:
            return cls._normalize_spaces(re.sub(r"\s*\|\s*", " ", cls._strip_invisible_chars(text)))
        return cls._strip_invisible_chars(str(text or "")).strip()

    @classmethod
    def _append_unique_sample(cls, samples: list[str], text: str, *, limit: int = MAX_MATCHED_SAMPLE_ITEMS) -> bool:
        """중복 없는 샘플만 제한 개수까지 누적한다."""
        normalized = cls._normalize_sample_compare_key(text)
        if not normalized:
            return False
        if any(cls._normalize_sample_compare_key(existing) == normalized for existing in samples):
            return False
        if len(samples) < limit:
            display_text = cls._normalize_heading_display_text(text) if cls._looks_like_generic_heading_text(text) else text
            samples.append(display_text)
        return True

    @staticmethod
    def _get_attr(node: Any, name: str, default: Any = None) -> Any:
        """namespace 가 섞인 XML 속성을 local name 기준으로 조회한다."""
        for key, value in getattr(node, "attrib", {}).items():
            if str(key).split("}")[-1] == name:
                return value
        return default

    @classmethod
    def _is_inside_table(cls, node: Any) -> bool:
        """현재 node 가 상위 table/tbl 내부에 속하는지 확인한다."""
        return any(cls._localname(parent) in {"tbl", "table"} for parent in node.iterancestors())

    @classmethod
    def _extract_paragraph_text(cls, node: Any) -> str:
        """표 밖 paragraph 블록에서 실제 표시 텍스트를 추출한다."""
        parts: list[str] = []
        for child in node.iter():
            if child is not node and cls._is_inside_table(child):
                continue
            name = cls._localname(child)
            if name == "t" and child.text:
                parts.append(child.text)
            elif name in {"lineBreak", "break"}:
                parts.append("\n")
            elif name == "tab":
                parts.append("\t")
            if child.tail:
                parts.append(child.tail)
        return cls._normalize_spaces("".join(parts))

    @classmethod
    def _extract_paragraph_text_in_cell(cls, node: Any) -> str:
        """table cell 내부 paragraph 에서 텍스트를 추출한다."""
        parts: list[str] = []
        for child in node.iter():
            name = cls._localname(child)
            if name == "t" and child.text:
                parts.append(child.text)
            elif name in {"lineBreak", "break"}:
                parts.append("\n")
            elif name == "tab":
                parts.append("\t")
            if child.tail:
                parts.append(child.tail)
        return cls._normalize_spaces("".join(parts))

    @classmethod
    def _extract_text_from_cell(cls, tc: Any) -> str:
        """table cell 의 여러 paragraph 텍스트를 1개 문자열로 합친다."""
        para_texts: list[str] = []
        for node in tc.iter():
            if cls._localname(node) != "p":
                continue
            text = cls._extract_paragraph_text_in_cell(node)
            if text:
                para_texts.append(text)
        if para_texts:
            return "\n".join(para_texts)
        return ""

    @classmethod
    def _is_spacer_column_value(cls, value: str) -> bool:
        normalized = cls._normalize_spaces(value)
        if not normalized:
            return True
        return bool(re.fullmatch(r"[↓⇓↑⇑→←↔↕\-–—~|│┃ ]+", normalized))

    @classmethod
    def _strip_spacer_columns(cls, grid: list[list[str]]) -> list[list[str]]:
        if not grid:
            return grid
        max_cols = max((len(row) for row in grid), default=0)
        if max_cols <= 1:
            return grid

        keep_indices: list[int] = []
        for col_idx in range(max_cols):
            column_values = [row[col_idx] if col_idx < len(row) else "" for row in grid]
            if all(cls._is_spacer_column_value(value) for value in column_values):
                continue
            keep_indices.append(col_idx)

        if len(keep_indices) == max_cols or not keep_indices:
            return grid
        return [
            [row[idx] if idx < len(row) else "" for idx in keep_indices]
            for row in grid
        ]

    @classmethod
    def _extract_table(cls, table_elem: Any) -> list[list[str]]:
        """rowSpan/colSpan 을 반영해 table XML 을 2차원 문자열 행렬로 변환한다."""
        grid: list[list[str]] = []
        row_spans: dict[int, dict[str, Any]] = {}
        trs = [node for node in table_elem.iterchildren() if cls._localname(node) == "tr"]

        for tr in trs:
            row: list[str] = []
            col_idx = 0
            tcs = [node for node in tr.iterchildren() if cls._localname(node) == "tc"]
            tc_idx = 0

            while tc_idx < len(tcs) or col_idx in row_spans:
                if col_idx in row_spans:
                    info = row_spans[col_idx]
                    row.append(info["text"])
                    info["remaining"] -= 1
                    if info["remaining"] <= 0:
                        del row_spans[col_idx]
                    col_idx += 1
                    continue

                if tc_idx >= len(tcs):
                    break

                tc = tcs[tc_idx]
                tc_idx += 1
                cell_text = cls._extract_text_from_cell(tc)
                col_span = cls._get_cell_span(tc, "colSpan", 1)
                row_span = cls._get_cell_span(tc, "rowSpan", 1)

                for offset in range(col_span):
                    curr_col = col_idx + offset
                    # 병합 셀은 첫 칸에만 텍스트를 두고 나머지는 빈칸으로 두어 비교를 단순화한다.
                    value = cell_text if offset == 0 else ""
                    row.append(value)
                    if row_span > 1:
                        row_spans[curr_col] = {"text": cell_text, "remaining": row_span - 1}
                col_idx += col_span

            if any((cell or "").strip() for cell in row):
                grid.append(row)

        return cls._strip_spacer_columns(grid)

    @classmethod
    def _get_cell_span(cls, tc: Any, name: str, default: int) -> int:
        """tc 또는 하위 cellSpan 노드에서 row/col span 값을 읽는다."""
        value = cls._get_attr(tc, name)
        try:
            if value is not None:
                return int(value)
        except (TypeError, ValueError):
            pass
        for child in tc.iterchildren():
            if cls._localname(child) == "cellSpan":
                nested_value = cls._get_attr(child, name)
                try:
                    if nested_value is not None:
                        return int(nested_value)
                except (TypeError, ValueError):
                    pass
        return default

    @classmethod
    def _is_table_wrapper_paragraph(cls, node: Any) -> bool:
        """table 만 감싸고 자체 텍스트는 없는 wrapper paragraph 인지 판정한다."""
        has_direct_table = any(cls._localname(child) in {"tbl", "table"} for child in node.iterdescendants())
        if not has_direct_table:
            return False
        visible_text = cls._extract_paragraph_text(node)
        return not bool(visible_text)

    @staticmethod
    def _extract_paragraph_text_from_xml(xml_text: str) -> str:
        """손상된 XML fallback 경로에서 paragraph 텍스트를 regex 기반으로 추출한다."""
        working = re.sub(r"<[^>]+:(lineBreak|break)\b[^>]*/?>", "\n", xml_text)
        working = re.sub(r"<[^>]+:tab\b[^>]*/?>", "\t", working)
        text_parts = re.findall(r"<[^>]+:t\b[^>]*>(.*?)</[^>]+:t>", working, flags=re.DOTALL)
        if text_parts:
            return HwpxHeadingParser._normalize_spaces(html.unescape("".join(text_parts)))
        stripped = re.sub(r"<[^>]+>", " ", working)
        return HwpxHeadingParser._normalize_spaces(html.unescape(stripped))

    @staticmethod
    def _bool_option(options: dict[str, Any], key: str, default: bool) -> bool:
        """rule option 을 bool 로 읽되 값이 없으면 기본값을 사용한다."""
        value = options.get(key)
        if value is None:
            return default
        return bool(value)

    @staticmethod
    def _int_option(options: dict[str, Any], key: str, default: int) -> int:
        """rule option 을 양수 int 로 읽는다."""
        value = options.get(key)
        try:
            parsed = int(value)
        except (TypeError, ValueError):
            return default
        return parsed if parsed > 0 else default

    @staticmethod
    def _nonneg_int_option(options: dict[str, Any], key: str) -> int | None:
        """rule option 을 0 이상 int 로 읽고 실패 시 None 을 반환한다."""
        value = options.get(key)
        if value is None:
            return None
        try:
            parsed = int(value)
        except (TypeError, ValueError):
            return None
        return parsed if parsed >= 0 else None

    @staticmethod
    def _get_rule_options(level: dict[str, Any]) -> dict[str, Any]:
        """level 에서 rule_options dict 만 안전하게 꺼낸다."""
        options = level.get("rule_options")
        return options if isinstance(options, dict) else {}

    @classmethod
    def _space_pattern(cls, max_spaces: int | None, *, default_unbounded: str) -> str:
        """공백 허용 범위를 정규식 fragment 로 만든다."""
        if max_spaces is None:
            return default_unbounded
        return rf"[ \t]{{0,{max_spaces}}}"

    @classmethod
    def _prefix_with_spacing(cls, options: dict[str, Any]) -> str:
        """marker 앞쪽 공백 허용 규칙을 정규식 prefix 로 만든다."""
        leading_space_max = cls._nonneg_int_option(options, "leading_space_max")
        if leading_space_max is not None:
            return r"^" + cls._space_pattern(leading_space_max, default_unbounded=r"[ \t]*")
        return r"^\s*" if cls._bool_option(options, "allow_leading_space", True) else r"^"

    @classmethod
    def _suffix_with_spacing(cls, options: dict[str, Any]) -> str:
        """marker 뒤쪽 공백/문장종료 허용 규칙을 정규식 suffix 로 만든다."""
        trailing_space_max = cls._nonneg_int_option(options, "trailing_space_max")
        if trailing_space_max is not None:
            spacing = cls._space_pattern(trailing_space_max, default_unbounded=r"[ \t]*")
            return spacing + r"$"
        return r"(?:\s+|$)" if cls._bool_option(options, "require_space_or_eol", True) else ""

    @classmethod
    def _title_text_suffix(cls, options: dict[str, Any]) -> str:
        """marker 뒤에 실제 제목 텍스트가 따라와야 하는 조건을 만든다."""
        trailing_space_max = cls._nonneg_int_option(options, "trailing_space_max")
        spacing = cls._space_pattern(trailing_space_max, default_unbounded=r"\s*")
        if cls._bool_option(options, "require_text_after_marker", True):
            return spacing + r"(?=.*[가-힣A-Za-z]).+$"
        return cls._suffix_with_spacing(options)

    @staticmethod
    def _trailing_dot_mode(options: dict[str, Any], notation: str, *, default_optional: bool = False) -> str:
        """숫자/문자 marker 뒤 점(.) 허용 정책을 결정한다."""
        explicit = str(options.get("trailing_dot") or "").strip().lower()
        if explicit in {"required", "forbidden", "either"}:
            return explicit
        if " / " in notation:
            return "either"
        if default_optional:
            return "either"
        return "required" if notation.strip().endswith(".") else "forbidden"

    @staticmethod
    def _roman_marker_pattern(roman_range: str, mode: str) -> str:
        """유니코드/ASCII 로마자 marker 를 함께 허용하는 정규식 fragment 를 만든다."""
        base = rf"(?:[{roman_range}]|[IVXLCDM]+)"
        if mode == "forbidden":
            return base
        if mode == "either":
            return base + r"(?:\.)?"
        return base + r"\."

    @classmethod
    def _compile_level_pattern(cls, level: dict[str, Any]) -> str:
        """rule_type 과 options 를 조합해 level 매칭용 정규식을 생성한다."""
        rule_type = str(level.get("rule_type") or "").strip()
        options = cls._get_rule_options(level)
        notation = str(level.get("notation") or "").strip()

        if rule_type == "numeric_dot":
            segments = cls._int_option(options, "segments", 1)
            if segments <= 1:
                mode = cls._trailing_dot_mode(
                    options,
                    notation,
                    default_optional=cls._bool_option(options, "allow_missing_terminal_dot", False),
                )
                body = r"[0-9]+" if mode == "forbidden" else r"[0-9]+(?:\.)?" if mode == "either" else r"[0-9]+\."
            else:
                body = r"\.".join([r"[0-9]+"] * segments)
                mode = cls._trailing_dot_mode(
                    options,
                    notation,
                    default_optional=cls._bool_option(options, "allow_trailing_dot", True),
                )
                if mode == "required":
                    body += r"\."
                elif mode == "either":
                    body += r"\.?"
            return cls._prefix_with_spacing(options) + body + cls._title_text_suffix(options)

        if rule_type == "korean_letter_dot":
            letter_range = str(options.get("letter_range") or "가-히")
            mode = cls._trailing_dot_mode(options, notation)
            body = f"[{letter_range}]" if mode == "forbidden" else f"[{letter_range}](?:\\.)?" if mode == "either" else f"[{letter_range}]\\."
            return cls._prefix_with_spacing(options) + body + cls._title_text_suffix(options)

        if rule_type == "roman":
            roman_range = str(options.get("roman_range") or "ⅠⅡⅢⅣⅤⅥⅦⅧⅨⅩ")
            mode = cls._trailing_dot_mode(options, notation, default_optional=True)
            body = cls._roman_marker_pattern(roman_range, mode)
            return cls._prefix_with_spacing(options) + body + cls._title_text_suffix(options)

        if rule_type == "numeric_paren":
            return cls._prefix_with_spacing(options) + r"[0-9]+\)" + cls._title_text_suffix(options)
        if rule_type == "korean_letter_paren":
            letter_range = str(options.get("letter_range") or "가-히")
            return cls._prefix_with_spacing(options) + f"[{letter_range}]\\)" + cls._title_text_suffix(options)
        if rule_type == "circled_number":
            return cls._prefix_with_spacing(options) + r"[\u2460-\u2473]" + cls._title_text_suffix(options)
        if rule_type == "circled_korean":
            return cls._prefix_with_spacing(options) + r"[\u3260-\u326F]" + cls._title_text_suffix(options)
        if rule_type == "paren_numeric":
            return cls._prefix_with_spacing(options) + r"\([0-9]+\)" + cls._title_text_suffix(options)
        if rule_type == "paren_korean":
            letter_range = str(options.get("letter_range") or "가-히")
            return cls._prefix_with_spacing(options) + f"\\([{letter_range}]\\)" + cls._title_text_suffix(options)
        if rule_type == "legal_article":
            body = r"제[0-9]+조"
            if cls._bool_option(options, "allow_sub_article", True):
                body += r"(?:의[0-9]+)?"
            if cls._bool_option(options, "allow_title_paren", True):
                body += r"(?:\s*\([^\)]+\))?"
            return cls._prefix_with_spacing(options) + body + cls._title_text_suffix(options)
        if rule_type == "symbol":
            symbols = options.get("symbols")
            text = "".join(str(symbol) for symbol in symbols if str(symbol)) if isinstance(symbols, list) else str(symbols or level.get("notation") or "")
            if text:
                return cls._prefix_with_spacing(options) + f"[{re.escape(text)}]" + cls._title_text_suffix(options)

        pattern = level.get("pattern")
        return str(pattern or "")

    @staticmethod
    def _normalize_level_location(value: Any) -> str:
        """schema level 의 location 값을 paragraph/table/both 로 정규화한다."""
        text = str(value or "").strip().lower()
        return text if text in {"paragraph", "table", "both"} else "paragraph"

    @staticmethod
    def _normalize_source_location(value: Any) -> str:
        """raw source location 값을 paragraph/table 2종으로 정규화한다."""
        text = str(value or "").strip().lower()
        return "table" if text in {"table", "table-cell"} else "paragraph"

    @classmethod
    def _location_matches(cls, level_location: str, source_location: str) -> bool:
        """schema level location 과 source line location 이 호환되는지 확인한다."""
        normalized_level = cls._normalize_level_location(level_location)
        normalized_source = cls._normalize_source_location(source_location)
        return normalized_level == "both" or normalized_level == normalized_source

    @classmethod
    def _strip_heading_marker(cls, level: dict[str, Any], line: str) -> str:
        """heading marker 를 제거한 뒤 제목 본문만 남긴다."""
        text = (line or "").lstrip()
        rule_type = str(level.get("rule_type") or "").strip()
        options = cls._get_rule_options(level)
        notation = str(level.get("notation") or "").strip()

        if rule_type == "numeric_dot":
            segments = cls._int_option(options, "segments", 1)
            if segments <= 1:
                mode = cls._trailing_dot_mode(options, notation, default_optional=cls._bool_option(options, "allow_missing_terminal_dot", False))
                marker = re.compile(r"^[0-9]+") if mode == "forbidden" else re.compile(r"^[0-9]+(?:\.)?") if mode == "either" else re.compile(r"^[0-9]+\.")
            else:
                body = r"\.".join([r"[0-9]+"] * segments)
                mode = cls._trailing_dot_mode(options, notation, default_optional=cls._bool_option(options, "allow_trailing_dot", True))
                marker = re.compile(rf"^{body}\.") if mode == "required" else re.compile(rf"^{body}\.?") if mode == "either" else re.compile(rf"^{body}")
        elif rule_type == "korean_letter_dot":
            letter_range = str(options.get("letter_range") or "가-히")
            mode = cls._trailing_dot_mode(options, notation)
            marker = re.compile(rf"^[{letter_range}]") if mode == "forbidden" else re.compile(rf"^[{letter_range}](?:\.)?") if mode == "either" else re.compile(rf"^[{letter_range}]\.")
        elif rule_type == "roman":
            roman_range = str(options.get("roman_range") or "ⅠⅡⅢⅣⅤⅥⅦⅧⅨⅩ")
            mode = cls._trailing_dot_mode(options, notation, default_optional=True)
            marker = re.compile(r"^" + cls._roman_marker_pattern(roman_range, mode))
        elif rule_type == "numeric_paren":
            marker = re.compile(r"^[0-9]+\)")
        elif rule_type == "korean_letter_paren":
            letter_range = str(options.get("letter_range") or "가-히")
            marker = re.compile(rf"^[{letter_range}]\)")
        elif rule_type == "circled_number":
            marker = re.compile(r"^[\u2460-\u2473]")
        elif rule_type == "circled_korean":
            marker = re.compile(r"^[\u3260-\u326F]")
        elif rule_type == "paren_numeric":
            marker = re.compile(r"^\([0-9]+\)")
        elif rule_type == "paren_korean":
            letter_range = str(options.get("letter_range") or "가-히")
            marker = re.compile(rf"^\([{letter_range}]\)")
        elif rule_type == "legal_article":
            marker = re.compile(r"^제[0-9]+조(?:의[0-9]+)?(?:\s*\([^\)]+\))?")
        else:
            return text

        match = marker.match(text)
        return text[match.end():].strip() if match else text

    @classmethod
    def _is_heading_like_line(cls, level: dict[str, Any], line: str) -> bool:
        """정규식 매치 후에도 실제 heading 문장처럼 보이는지 추가 판정한다."""
        rule_type = str(level.get("rule_type") or "").strip()
        if rule_type not in {"numeric_dot", "korean_letter_dot", "roman", "numeric_paren", "korean_letter_paren", "legal_article"}:
            return True
        remainder = cls._strip_heading_marker(level, line)
        return bool(remainder and not remainder[0].isdigit() and len(remainder) <= 80 and re.search(r"[가-힣A-Za-z]", remainder))

    @staticmethod
    def _is_heading_like_symbol_candidate(text: str) -> bool:
        """신규 symbol heading 후보로 볼 수 있는 텍스트인지 판정한다."""
        line = (text or "").strip()
        match = _SYMBOL_CANDIDATE_RE.match(line)
        if not match:
            return False
        remainder = line[match.end():].strip()
        return bool(remainder and not remainder[0].isdigit() and len(remainder) <= 80 and re.search(r"[가-힣A-Za-z]", remainder))

    @classmethod
    def _fallback_match_existing(
        cls,
        text: str,
        location: str,
        *,
        is_box_table: bool,
        existing_matchers: list[tuple[dict[str, Any], str, re.Pattern[str]]],
    ) -> tuple[dict[str, Any], str, re.Pattern[str]] | None:
        """정규식이 조금 달라도 rule_type 이 같은 기존 level 로 보정 매칭한다."""
        line = (text or "").strip()
        if not line:
            return None
        checks = [
            (r"^\s*(?:[ⅠⅡⅢⅣⅤⅥⅦⅧⅨⅩ]|[IVXLCDM]+)\s*[\.．。]?\s+(.+)$", "roman"),
            (r"^\s*[가-힣]\s*[\.．。]\s*(.+)$", "korean_letter_dot"),
            (r"^\s*[가-힣]\s*\)\s*(.+)$", "korean_letter_paren"),
            (r"^\s*[0-9]+\s*[\.．。]\s+(.+)$", "numeric_dot"),
            (r"^\s*[0-9]+\s*\)\s*(.+)$", "numeric_paren"),
        ]
        for pattern, rule_type in checks:
            match = re.match(pattern, line)
            if not match:
                continue
            remainder = (match.group(1) or "").strip()
            if not remainder or not re.search(r"[가-힣A-Za-z]", remainder) or remainder[0].isdigit():
                continue
            for level, row_id, compiled in existing_matchers:
                if is_box_table and cls._normalize_level_location(level.get("location") or "paragraph") == "paragraph":
                    continue
                if not cls._location_matches(str(level.get("location") or "paragraph"), location):
                    continue
                if str(level.get("rule_type") or "").strip() == rule_type:
                    return level, row_id, compiled
        return None

    @classmethod
    def _fallback_match_existing_symbol_only(
        cls,
        text: str,
        location: str,
        *,
        is_box_table: bool,
        existing_matchers: list[tuple[dict[str, Any], str, re.Pattern[str]]],
    ) -> tuple[dict[str, Any], str, re.Pattern[str]] | None:
        """symbol 계열 marker 는 rule_type/allowed symbol 기준으로 느슨하게 재매칭한다."""
        line = (text or "").strip()
        match = _SYMBOL_CANDIDATE_RE.match(line)
        if not match:
            return None
        symbol = match.group(1)
        for level, row_id, compiled in existing_matchers:
            if is_box_table and cls._normalize_level_location(level.get("location") or "paragraph") == "paragraph":
                continue
            if not cls._location_matches(str(level.get("location") or "paragraph"), location):
                continue
            rule_type = str(level.get("rule_type") or "").strip()
            notation = str(level.get("notation") or "").strip()
            rule_options = cls._get_rule_options(level)
            if rule_type == "circled_number" and re.match(r"[①-⑳]", symbol):
                return level, row_id, compiled
            if rule_type == "circled_korean" and re.match(r"[㉮-㉻]", symbol):
                return level, row_id, compiled
            if rule_type == "symbol":
                symbols = rule_options.get("symbols")
                allowed_symbols = {str(value) for value in symbols if str(value)} if isinstance(symbols, list) else ({notation} if notation else set())
                if symbol in allowed_symbols:
                    return level, row_id, compiled
        return None

    @classmethod
    def _matches_appendix_title_table(cls, level: dict[str, Any], item: dict[str, Any]) -> bool:
        """appendix 시작용 표 제목 규칙과 raw table line 이 맞는지 판정한다."""
        source_location = cls._normalize_source_location(item.get("location") or "paragraph")
        if source_location != "table":
            return False
        options = cls._get_rule_options(level)
        text = cls._normalize_single_row_table_heading_text(item, str(item.get("text") or ""))
        if not text:
            return False
        row_count = item.get("row_count")
        cell_count = item.get("cell_count")
        if cls._bool_option(options, "single_row_table_only", True) and row_count not in {None, 1}:
            return False
        required_cell_count = options.get("require_cell_count")
        try:
            expected_cell_count = int(required_cell_count) if required_cell_count is not None else None
        except (TypeError, ValueError):
            expected_cell_count = None
        if expected_cell_count is not None and cell_count not in {None, expected_cell_count}:
            return False
        keywords = options.get("keywords")
        keyword_group = (
            "|".join(re.escape(str(keyword).strip()) for keyword in keywords if str(keyword).strip())
            if isinstance(keywords, list)
            else ""
        ) or r"참고|첨부"
        title_min_length = cls._int_option(options, "title_cell_min_length", 2)
        pattern = (
            rf"^\s*(?:{keyword_group})\s*\d+\s*(?:\||\s)\s*(.+)$"
            if cls._bool_option(options, "left_cell_number_required", True)
            else rf"^\s*(?:{keyword_group})\s*(?:\||\s)\s*(.+)$"
        )
        match = re.match(pattern, text)
        if not match:
            return False
        title_text = cls._normalize_spaces(match.group(1) or "")
        return len(title_text) >= title_min_length

    @classmethod
    def _matches_level_for_preparse(
        cls,
        level: dict[str, Any],
        item: dict[str, Any],
        *,
        restrict_box_main_reentry: bool = False,
    ) -> bool:
        """raw line 1건이 특정 schema level 규칙에 매칭되는지 판정한다."""
        source_location = cls._normalize_source_location(item.get("location") or "paragraph")
        level_location = str(level.get("location") or "paragraph")
        if restrict_box_main_reentry and bool(item.get("is_box")) and source_location == "table":
            if cls._normalize_level_location(level_location) == "paragraph":
                return False
        if not cls._location_matches(level_location, source_location):
            return False
        text = cls._normalize_single_row_table_heading_text(item, str(item.get("text") or ""))
        if not text:
            return False
        if str(level.get("rule_type") or "").strip() == "appendix_title_table":
            return cls._matches_appendix_title_table(level, item)
        pattern = cls._compile_level_pattern(level)
        if not pattern:
            return False
        try:
            compiled = re.compile(pattern)
        except re.error:
            return False
        return bool(compiled.match(text) and cls._is_heading_like_line(level, text))

    @staticmethod
    def _is_symbolic_reentry_level(level: dict[str, Any]) -> bool:
        """body 복귀 기준에서 제외할 상징형 level 인지 판정한다."""
        rule_type = str(level.get("rule_type") or "").strip().lower()
        rule_id = str(level.get("rule_id") or "").strip().lower()
        notation = str(level.get("notation") or "").strip()
        if rule_type in {"symbol", "custom:symbol"} or rule_id in {"symbol", "custom:symbol"}:
            return True
        compact = re.sub(r"[\s,.\-_/|()]+", "", notation)
        return bool(compact and not re.search(r"[0-9A-Za-z가-힣ⅠⅡⅢⅣⅤⅥⅦⅧⅨⅩ①-⑳㉮-㉻]", compact))

    @staticmethod
    def _level_depth(level: dict[str, Any], fallback_idx: int = 0) -> int:
        """level dict 에서 depth 를 안전하게 읽는다."""
        return int(level.get("depth") or fallback_idx or 0)

    @classmethod
    def _find_first_matching_level(
        cls,
        item: dict[str, Any],
        schema: dict[str, Any] | None,
        *,
        depth: int | None = None,
        restrict_box_main_reentry: bool = False,
        exclude_symbolic_reentry_levels: bool = False,
    ) -> dict[str, Any] | None:
        """schema 에서 item 과 처음 매칭되는 level 을 찾는다."""
        if not isinstance(schema, dict):
            return None
        levels = schema.get("levels")
        if not isinstance(levels, list):
            return None
        for idx, raw_level in enumerate(levels, start=1):
            if not isinstance(raw_level, dict):
                continue
            level_depth = cls._level_depth(raw_level, idx)
            if depth is not None and level_depth != depth:
                continue
            if exclude_symbolic_reentry_levels and cls._is_symbolic_reentry_level(raw_level):
                continue
            if cls._matches_level_for_preparse(raw_level, item, restrict_box_main_reentry=restrict_box_main_reentry):
                return raw_level
        return None

    @classmethod
    def _iter_matching_levels_for_exit_criteria(
        cls,
        criteria: dict[str, Any] | None,
        target_schema: dict[str, Any] | None,
    ) -> list[dict[str, Any]]:
        """exit criteria 가 가리키는 target schema level 목록을 추린다.

        depth 가 지정되면 해당 depth 와 그 상위 depth(더 작은 depth)도 함께 포함한다.
        """
        if not isinstance(criteria, dict) or not isinstance(target_schema, dict):
            return []
        if str(criteria.get("mode") or "").strip() != "matched_heading":
            return []
        match = criteria.get("match")
        if not isinstance(match, dict):
            return []
        levels = target_schema.get("levels")
        if not isinstance(levels, list):
            return []
        target_depth = match.get("depth")
        target_rule_id = str(match.get("rule_id") or "").strip()
        target_notation = str(match.get("notation") or "").strip()
        matched_levels: list[dict[str, Any]] = []
        for idx, level in enumerate(levels, start=1):
            if not isinstance(level, dict):
                continue
            level_depth = int(level.get("depth") or idx)
            if target_depth is not None:
                try:
                    if level_depth > int(target_depth):
                        continue
                except (TypeError, ValueError):
                    continue
            if target_rule_id and str(level.get("rule_id") or "").strip() != target_rule_id:
                continue
            if target_notation and str(level.get("notation") or "").strip() != target_notation:
                continue
            matched_levels.append(level)
        return matched_levels

    @classmethod
    def _matches_exit_criteria(
        cls,
        item: dict[str, Any],
        *,
        criteria: dict[str, Any] | None,
        target_schema: dict[str, Any] | None,
        restrict_box_main_reentry: bool = False,
        exclude_symbolic_reentry_levels: bool = False,
    ) -> bool:
        """현재 raw line 이 명시적 exit criteria 를 만족하는지 판정한다."""
        for level in cls._iter_matching_levels_for_exit_criteria(criteria, target_schema):
            if exclude_symbolic_reentry_levels and cls._is_symbolic_reentry_level(level):
                continue
            if cls._matches_level_for_preparse(level, item, restrict_box_main_reentry=restrict_box_main_reentry):
                return True
        return False

    @classmethod
    def _matches_default_appendix_exit(cls, item: dict[str, Any], main_schema: dict[str, Any]) -> bool:
        """appendix 종료 기준이 없을 때 기본 body 복귀 규칙을 적용한다."""
        levels = main_schema.get("levels")
        if not isinstance(levels, list):
            return False
        for idx, level in enumerate(levels, start=1):
            if not isinstance(level, dict):
                continue
            if int(level.get("depth") or idx) > 2:
                continue
            if cls._is_symbolic_reentry_level(level):
                continue
            if cls._matches_level_for_preparse(level, item, restrict_box_main_reentry=True):
                return True
        return False

    @classmethod
    def _should_skip_main_preparse_item(cls, item: dict[str, Any]) -> bool:
        """body 후보 계산에서 제외할 table line 인지 판정한다."""
        source_location = cls._normalize_source_location(item.get("location") or "paragraph")
        if source_location != "table":
            return False
        if bool(item.get("is_box")):
            return True
        row_count = item.get("row_count")
        cell_count = item.get("cell_count")
        if row_count == 1 and isinstance(cell_count, int) and cell_count > 1:
            text = str(item.get("text") or "").strip()
            # 1x2 배너형 표는 본문 depth 1 heading 으로 쓰이는 경우가 많아
            # heading marker 가 보이면 skip 하지 않고 기존 규칙 매칭까지 보낸다.
            return not cls._looks_like_generic_heading_text(text)
        return False

    @classmethod
    def _is_single_row_table_item(cls, item: dict[str, Any]) -> bool:
        """raw line 이 1행 table 기반 항목인지 판정한다."""
        source_location = cls._normalize_source_location(item.get("location") or "paragraph")
        if source_location != "table":
            return False
        row_count = item.get("row_count")
        return row_count == 1


MAX_DOCUMENT_BYTES = 30 * 1024 * 1024
SUPPORTED_DOCUMENT_SUFFIXES = (".pdf", ".docx", ".hwpx")


def _normalize_document_paragraph(text: str) -> str:
    """PDF/DOCX/HWPX 공통 paragraph 정규화."""
    if not text:
        return ""
    lines = [" ".join(line.split()) for line in text.splitlines()]
    return "\n".join(line for line in lines if line.strip()).strip()


def _is_document_heading(text: str) -> bool:
    """문서 ingest 관점에서 heading 으로 취급할 패턴을 판정한다."""
    t = (text or "").strip()
    if not t or len(t) > 120:
        return False
    patterns = (
        r"^[ⅠⅡⅢⅣⅤⅥⅦⅧⅨⅩ]\s+.+",
        r"^[IVX]+\s+.+",
        r"^제\s*\d+\s*[장절관조]\b",
        r"^\d+\.\s+\S",
        r"^\d+\)\s+\S",
        r"^[가-힣]\.\s+",
        r"^[가나다라마바사아자차카타파하]\.\s+",
        r"^[①②③④⑤⑥⑦⑧⑨⑩⑪⑫⑬⑭⑮⑯⑰⑱⑲⑳]\s*",
        r"^[□■◇◆○●▣▶◀]\s*",
    )
    return any(re.match(pattern, t) for pattern in patterns)


def _extract_pdf_blocks(content: bytes, filename: str) -> list[ParsedBlock]:
    """PDF bytes 를 ParsedBlock 목록으로 변환한다."""
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise HwpxParsingError(
            "pypdf 모듈이 설치되지 않았습니다. requirements.txt 의 pypdf 를 설치해 주세요."
        ) from exc

    try:
        reader = PdfReader(io.BytesIO(content))
    except Exception as exc:  # noqa: BLE001
        raise HwpxParsingError(f"PDF 를 열 수 없습니다: {exc}") from exc

    blocks: list[ParsedBlock] = []
    current_heading = ""
    for page_idx, page in enumerate(reader.pages, start=1):
        try:
            raw = page.extract_text() or ""
        except Exception:  # noqa: BLE001
            continue
        for para in re.split(r"\n{2,}", raw):
            text = _normalize_document_paragraph(para)
            if not text:
                continue
            if _is_document_heading(text):
                current_heading = text
                continue
            blocks.append(ParsedBlock(title=current_heading, text=text, section=f"page_{page_idx}"))
    if not blocks:
        raise HwpxParsingError(f"{filename} 에서 파싱 가능한 텍스트를 찾지 못했어요.")
    return blocks


def _extract_docx_blocks(content: bytes, filename: str) -> list[ParsedBlock]:
    """DOCX bytes 를 ParsedBlock 목록으로 변환한다."""
    try:
        from docx import Document  # type: ignore[import-not-found]
    except ImportError as exc:
        raise HwpxParsingError(
            "python-docx 모듈이 설치되지 않았습니다. requirements.txt 의 python-docx 를 설치해 주세요."
        ) from exc

    try:
        doc = Document(io.BytesIO(content))
    except Exception as exc:  # noqa: BLE001
        raise HwpxParsingError(f"DOCX 를 열 수 없습니다: {exc}") from exc

    blocks: list[ParsedBlock] = []
    current_heading = ""
    for para in doc.paragraphs:
        text = _normalize_document_paragraph(para.text or "")
        if not text:
            continue
        style_name = (para.style.name if para.style else "") or ""
        if style_name.lower().startswith("heading") or _is_document_heading(text):
            current_heading = text
            continue
        blocks.append(ParsedBlock(title=current_heading, text=text, section=""))

    for tbl in doc.tables:
        for row in tbl.rows:
            row_text_parts = [_normalize_document_paragraph(cell.text or "") for cell in row.cells]
            row_text = " | ".join(part for part in row_text_parts if part)
            if row_text:
                blocks.append(ParsedBlock(title=current_heading, text=row_text, section="table"))

    if not blocks:
        raise HwpxParsingError(f"{filename} 에서 파싱 가능한 텍스트를 찾지 못했어요.")
    return blocks


def _extract_hwpx_blocks(content: bytes, filename: str) -> list[ParsedBlock]:
    """HWPX bytes 를 paragraph 기반 ParsedBlock 목록으로 변환한다."""
    with tempfile.NamedTemporaryFile(suffix=".hwpx", delete=False) as tmp:
        tmp.write(content)
        tmp_path = Path(tmp.name)
    try:
        parser = HwpxHeadingParser(main_schema={"levels": []})
        blocks: list[ParsedBlock] = []
        current_heading = ""
        with zipfile.ZipFile(tmp_path) as archive:
            section_files = sorted(
                [name for name in archive.namelist() if re.match(r"^Contents/section\d+\.xml$", name)],
                key=HwpxHeadingParser._sort_section_key,
            )
            if not section_files:
                raise HwpxParsingError("HWPX section XML을 찾지 못했습니다.")
            for section_name in section_files:
                xml_text = archive.read(section_name).decode("utf-8", errors="ignore")
                paragraph_blocks = re.findall(r"<[^>]+:p\b.*?</[^>]+:p>", xml_text, flags=re.DOTALL)
                for paragraph_xml in paragraph_blocks:
                    text = parser._extract_paragraph_text_from_xml(paragraph_xml)
                    if not text:
                        continue
                    if _is_document_heading(text):
                        current_heading = text
                        continue
                    blocks.append(ParsedBlock(title=current_heading, text=text, section=section_name))
    except HwpxParsingError:
        raise
    except Exception as exc:  # noqa: BLE001
        raise HwpxParsingError(f"HWPX 를 파싱할 수 없습니다: {exc}") from exc
    finally:
        try:
            tmp_path.unlink(missing_ok=True)
        except Exception:  # noqa: BLE001
            pass

    if not blocks:
        raise HwpxParsingError(f"{filename} 에서 파싱 가능한 텍스트를 찾지 못했어요.")
    return blocks


def extract_blocks(filename: str, content: bytes) -> list[ParsedBlock]:
    """파일 확장자에 따라 적절한 문서 파서를 호출해 ParsedBlock 목록을 생성한다."""
    if len(content) > MAX_DOCUMENT_BYTES:
        raise HwpxParsingError(
            f"파일이 너무 큽니다 ({len(content) // 1024 // 1024}MB). 30MB 이하만 업로드해 주세요."
        )
    suffix = Path(filename).suffix.lower()
    if suffix == ".pdf":
        return _extract_pdf_blocks(content, filename)
    if suffix == ".docx":
        return _extract_docx_blocks(content, filename)
    if suffix == ".hwpx":
        return _extract_hwpx_blocks(content, filename)
    raise HwpxParsingError(
        f"지원하지 않는 형식입니다: {suffix}. {', '.join(SUPPORTED_DOCUMENT_SUFFIXES)} 만 가능합니다."
    )


def iter_raw_block_lines(blocks: list[ParsedBlock]) -> list[dict[str, str]]:
    """ParsedBlock 목록을 화면 pre-parse 용 raw line 목록으로 변환한다."""
    lines: list[dict[str, str]] = []
    seen: set[tuple[str, str]] = set()

    def append_once(value: str, *, location: str = "paragraph") -> None:
        normalized = HwpxHeadingParser._normalize_spaces(value)
        normalized_location = HwpxHeadingParser._normalize_source_location(location)
        key = (normalized, normalized_location)
        if not normalized or key in seen:
            return
        seen.add(key)
        lines.append({"text": normalized, "location": normalized_location})

    for block in blocks:
        title = (block.title or "").strip()
        if title:
            append_once(title, location="paragraph")
        text = block.text or ""
        for raw_line in text.splitlines():
            line = raw_line.strip()
            if line:
                append_once(line, location="paragraph")
    return lines


def preview_preparse_lines(lines: list[dict[str, Any]], limit: int = 100) -> list[dict[str, Any]]:
    """화면 preview 카드에 맞는 pre-parse line 요약을 만든다."""
    items: list[dict[str, Any]] = []
    for item in lines[:limit]:
        location = str(item.get("location") or "paragraph").strip()
        block_type = str(item.get("block_type") or location).strip()
        section_name = str(item.get("section_name") or "preparse").strip()
        items.append(
            {
                "title": f"[{block_type}]",
                "text": str(item.get("text") or ""),
                "section": f"{section_name}:{location}",
            }
        )
    return items


def split_preparse_lines_by_domain(
    source_lines: list[dict[str, Any]],
    *,
    main_schema: dict[str, Any],
    appendix_schema: dict[str, Any] | None,
    body_exit_criteria: dict[str, Any] | None = None,
    appendix_exit_criteria: dict[str, Any] | None = None,
) -> tuple[list[dict[str, Any]], list[dict[str, Any]]]:
    """본문/부록 schema 기준으로 raw line 을 분리한다."""
    parser = HwpxHeadingParser(
        main_schema=main_schema,
        appendix_schema=appendix_schema,
        body_exit_criteria=body_exit_criteria,
        appendix_exit_criteria=appendix_exit_criteria,
    )
    return parser._split_preparse_lines_by_domain(source_lines)


def build_heading_rows(
    schema: dict[str, Any] | None,
    *,
    raw_lines: list[dict[str, Any]] | None = None,
    allow_new_candidates: bool = True,
) -> list[dict[str, Any]]:
    """schema 와 raw line 으로 관리자 화면용 heading row 목록을 만든다."""
    parser = HwpxHeadingParser(main_schema=schema or {"levels": []})
    return parser._build_heading_rows(
        schema,
        raw_lines=raw_lines,
        allow_new_candidates=allow_new_candidates,
    )
