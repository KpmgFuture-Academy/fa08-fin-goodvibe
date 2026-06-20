from __future__ import annotations

import html
import io
import re
import zipfile
from enum import Enum
from pathlib import Path
from typing import Any
from xml.etree import ElementTree as ET

from .document_models import HeadingNode, ParsedSegment, RagEmbeddingParseError


class TablePattern(str, Enum):
    """
    HWPX 표를 RAG 문장화할 때 사용하는 대표 패턴 3종.

    1. HEADER_ROW_RECORDS
       - 첫째 행이 헤더, 나머지 행이 데이터인 일반 표
    2. HEADER_VALUE_PAIRS
       - 헤더-값 쌍이 가로로 반복되는 표
    3. HEADER_VALUE_PAIR_GROUPS
       - 2의 변형으로, 헤더-(소헤더:값) 목록 형태를 담는 표
    """

    HEADER_ROW_RECORDS = "header-row-records"
    HEADER_VALUE_PAIRS = "header-value-pairs"
    HEADER_VALUE_PAIR_GROUPS = "header-value-pair-groups"


class TableRowRole(str, Enum):
    """표 내부 개별 행의 역할."""

    EMPTY = "empty"
    HEADER_ROW = "header-row"
    DATA_ROW = "data-row"
    HEADER_VALUE_ROW = "header-value-row"


def _normalize_text(text: str) -> str:
    lines = [" ".join((line or "").split()) for line in (text or "").splitlines()]
    return "\n".join(line for line in lines if line.strip()).strip()


def _localname(tag: str) -> str:
    return tag.split("}", 1)[-1] if "}" in tag else tag


def _get_attr(node: ET.Element, name: str, default: Any = None) -> Any:
    for key, value in getattr(node, "attrib", {}).items():
        if str(key).split("}")[-1] == name:
            return value
    return default


def _normalize_schema(schema: dict[str, Any]) -> dict[str, Any]:
    levels = schema.get("levels")
    if not isinstance(levels, list):
        return {"levels": []}
    normalized: list[dict[str, Any]] = []
    for level in levels:
        if not isinstance(level, dict):
            continue
        next_level = dict(level)
        pattern = next_level.get("pattern")
        if isinstance(pattern, str) and "\\\\" in pattern:
            next_level["pattern"] = pattern.replace("\\\\", "\\")
        normalized.append(next_level)
    return {**schema, "levels": normalized}


def _normalize_level_location(value: Any) -> str:
    text = str(value or "").strip().lower()
    if text in {"paragraph", "table", "both"}:
        return text
    return "paragraph"


def _normalize_event_location(value: Any) -> str:
    text = str(value or "").strip().lower()
    if text in {"table", "table-cell"}:
        return "table"
    return "paragraph"


def _location_matches(level_location: str, event_location: str) -> bool:
    normalized_level = _normalize_level_location(level_location)
    normalized_event = _normalize_event_location(event_location)
    return normalized_level == "both" or normalized_level == normalized_event


def _bool_option(options: dict[str, Any], key: str, default: bool) -> bool:
    value = options.get(key)
    if value is None:
        return default
    return bool(value)


def _int_option(options: dict[str, Any], key: str, default: int) -> int:
    value = options.get(key)
    try:
        parsed = int(value)
    except (TypeError, ValueError):
        return default
    return parsed if parsed > 0 else default


def _nonneg_int_option(options: dict[str, Any], key: str) -> int | None:
    value = options.get(key)
    if value is None:
        return None
    try:
        parsed = int(value)
    except (TypeError, ValueError):
        return None
    return parsed if parsed >= 0 else None


def _get_rule_options(level: dict[str, Any]) -> dict[str, Any]:
    options = level.get("rule_options")
    return options if isinstance(options, dict) else {}


def _space_pattern(max_spaces: int | None, *, default_unbounded: str) -> str:
    if max_spaces is None:
        return default_unbounded
    return rf"[ \t]{{0,{max_spaces}}}"


def _prefix_with_spacing(options: dict[str, Any]) -> str:
    leading_space_max = _nonneg_int_option(options, "leading_space_max")
    if leading_space_max is not None:
        return r"^" + _space_pattern(leading_space_max, default_unbounded=r"[ \t]*")
    return r"^\s*" if _bool_option(options, "allow_leading_space", True) else r"^"


def _suffix_with_spacing(options: dict[str, Any]) -> str:
    trailing_space_max = _nonneg_int_option(options, "trailing_space_max")
    if trailing_space_max is not None:
        spacing = _space_pattern(trailing_space_max, default_unbounded=r"[ \t]*")
        return spacing + r"$"
    return r"(?:\s+|$)" if _bool_option(options, "require_space_or_eol", True) else ""


def _title_text_suffix(options: dict[str, Any]) -> str:
    trailing_space_max = _nonneg_int_option(options, "trailing_space_max")
    spacing = _space_pattern(trailing_space_max, default_unbounded=r"\s*")
    if _bool_option(options, "require_text_after_marker", True):
        return spacing + r"(?=.*[가-힣A-Za-z]).+$"
    return _suffix_with_spacing(options)


def _trailing_dot_mode(options: dict[str, Any], notation: str, *, default_optional: bool = False) -> str:
    explicit = str(options.get("trailing_dot") or "").strip().lower()
    if explicit in {"required", "forbidden", "either"}:
        return explicit
    if " / " in notation:
        return "either"
    if default_optional:
        return "either"
    return "required" if notation.strip().endswith(".") else "forbidden"


def _compile_level_pattern(level: dict[str, Any]) -> str:
    rule_type = str(level.get("rule_type") or "").strip()
    options = _get_rule_options(level)
    notation = str(level.get("notation") or "").strip()

    if rule_type == "appendix_title_table":
        return ""

    if rule_type == "numeric_dot":
        segments = _int_option(options, "segments", 1)
        if segments <= 1:
            mode = _trailing_dot_mode(
                options,
                notation,
                default_optional=_bool_option(options, "allow_missing_terminal_dot", False),
            )
            if mode == "forbidden":
                body = r"[0-9]+"
            elif mode == "either":
                body = r"[0-9]+(?:\.)?"
            else:
                body = r"[0-9]+\."
        else:
            body = r"\.".join([r"[0-9]+"] * segments)
            mode = _trailing_dot_mode(
                options,
                notation,
                default_optional=_bool_option(options, "allow_trailing_dot", True),
            )
            if mode == "required":
                body += r"\."
            elif mode == "either":
                body += r"\.?"
        return _prefix_with_spacing(options) + body + _title_text_suffix(options)

    if rule_type == "korean_letter_dot":
        letter_range = str(options.get("letter_range") or "가-히")
        mode = _trailing_dot_mode(options, notation)
        if mode == "forbidden":
            body = f"[{letter_range}]"
        elif mode == "either":
            body = f"[{letter_range}](?:\\.)?"
        else:
            body = f"[{letter_range}]\\."
        return _prefix_with_spacing(options) + body + _title_text_suffix(options)

    if rule_type == "roman":
        roman_range = str(options.get("roman_range") or "ⅠⅡⅢⅣⅤⅥⅦⅧⅨⅩ")
        mode = _trailing_dot_mode(options, notation, default_optional=True)
        if mode == "forbidden":
            body = f"[{roman_range}]"
        elif mode == "either":
            body = f"[{roman_range}](?:\\.)?"
        else:
            body = f"[{roman_range}]\\."
        return _prefix_with_spacing(options) + body + _title_text_suffix(options)

    if rule_type == "numeric_paren":
        return _prefix_with_spacing(options) + r"[0-9]+\)" + _title_text_suffix(options)

    if rule_type == "korean_letter_paren":
        letter_range = str(options.get("letter_range") or "가-히")
        return _prefix_with_spacing(options) + f"[{letter_range}]\\)" + _title_text_suffix(options)

    if rule_type == "circled_number":
        return _prefix_with_spacing(options) + r"[\u2460-\u2473]" + _title_text_suffix(options)

    if rule_type == "circled_korean":
        return _prefix_with_spacing(options) + r"[\u3260-\u326F]" + _title_text_suffix(options)

    if rule_type == "paren_numeric":
        return _prefix_with_spacing(options) + r"\([0-9]+\)" + _title_text_suffix(options)

    if rule_type == "paren_korean":
        letter_range = str(options.get("letter_range") or "가-히")
        return _prefix_with_spacing(options) + f"\\([{letter_range}]\\)" + _title_text_suffix(options)

    if rule_type == "legal_article":
        body = r"제[0-9]+조"
        if _bool_option(options, "allow_sub_article", True):
            body += r"(?:의[0-9]+)?"
        if _bool_option(options, "allow_title_paren", True):
            body += r"(?:\s*\([^\)]+\))?"
        return _prefix_with_spacing(options) + body + _title_text_suffix(options)

    if rule_type == "symbol":
        symbols = options.get("symbols")
        if isinstance(symbols, list):
            text = "".join(str(symbol) for symbol in symbols if str(symbol))
        else:
            text = str(symbols or level.get("notation") or "")
        if text:
            return _prefix_with_spacing(options) + f"[{re.escape(text)}]" + _title_text_suffix(options)

    pattern = level.get("pattern")
    return str(pattern or "").strip()


def _extract_markdown_heading(line: str) -> tuple[int, str] | None:
    match = re.match(r"^(#{1,6})\s+(.+)$", line)
    if not match:
        return None
    return len(match.group(1)), _normalize_text(match.group(2))


def _extract_paragraph_text(node: ET.Element) -> str:
    parts: list[str] = []
    for child in node.iter():
        name = _localname(child.tag)
        if name == "t" and child.text:
            parts.append(child.text)
        elif name in {"lineBreak", "break"}:
            parts.append("\n")
        elif name == "tab":
            parts.append("\t")
        if child is not node and child.tail and child.tail.strip():
            parts.append(child.tail)
    return _normalize_text("".join(parts))


def _count_line_segments(node: ET.Element) -> int:
    return len(_get_line_segment_positions(node))


def _get_line_segment_positions(node: ET.Element) -> list[int]:
    positions: list[int] = []
    for child in node.iter():
        if _localname(child.tag).lower() != "linesegarray":
            continue
        for item in child:
            if _localname(item.tag).lower() != "lineseg":
                continue
            raw = item.attrib.get("textpos")
            try:
                value = int(str(raw or "0"))
            except (TypeError, ValueError):
                value = 0
            positions.append(max(0, value))
        if positions:
            break
    deduped: list[int] = []
    for value in positions:
        if deduped and deduped[-1] == value:
            continue
        deduped.append(value)
    return deduped


def _has_explicit_line_break(node: ET.Element) -> bool:
    return any(_localname(child.tag) in {"lineBreak", "break"} for child in node.iter())


def _split_visual_lines(text: str, line_count: int, positions: list[int] | None = None) -> str:
    normalized = _normalize_text(text)
    if line_count <= 1 or "\n" in normalized or len(normalized) < 12:
        return normalized

    if positions and len(positions) >= 2:
        valid_positions = [pos for pos in positions if 0 <= pos < len(normalized)]
        if valid_positions and valid_positions[0] != 0:
            valid_positions.insert(0, 0)
        elif not valid_positions:
            valid_positions = [0]
        if valid_positions[-1] != len(normalized):
            valid_positions.append(len(normalized))

        sliced_lines: list[str] = []
        for idx in range(len(valid_positions) - 1):
            start = valid_positions[idx]
            end = valid_positions[idx + 1]
            piece = _normalize_text(normalized[start:end])
            if piece:
                sliced_lines.append(piece)
        if sliced_lines:
            return "\n".join(sliced_lines)

    words = [part for part in normalized.split(" ") if part]
    if len(words) <= 1:
        return normalized

    target_count = min(line_count, len(words))
    total_length = len(normalized)
    target_lengths = [round(total_length * idx / target_count) for idx in range(1, target_count)]
    preferred_break_chars = set(",.;:)]}>") | {"'", '"'}

    lines: list[str] = []
    current_words: list[str] = []
    current_length = 0
    consumed_length = 0
    target_index = 0

    for idx, word in enumerate(words):
        projected = current_length + (1 if current_words else 0) + len(word)
        current_words.append(word)
        current_length = projected

        remaining_words = len(words) - idx - 1
        remaining_lines = target_count - len(lines) - 1
        if remaining_lines <= 0:
            continue
        if remaining_words < remaining_lines:
            continue

        should_break = False
        if target_index < len(target_lengths) and consumed_length + current_length >= target_lengths[target_index]:
            should_break = True
        joined = " ".join(current_words)
        if not should_break and word and word[-1] in preferred_break_chars and len(joined) >= 10:
            should_break = True

        if should_break:
            lines.append(joined)
            consumed_length += len(joined) + 1
            current_words = []
            current_length = 0
            target_index += 1

    if current_words:
        lines.append(" ".join(current_words))

    if len(lines) <= 1:
        return normalized
    return "\n".join(_normalize_text(line) for line in lines if _normalize_text(line))


def _extract_paragraph_text_excluding_tables(node: ET.Element) -> str:
    parts: list[str] = []

    def walk(current: ET.Element, *, inside_table: bool = False) -> None:
        name = _localname(current.tag)
        next_inside_table = inside_table or name == "tbl"

        if not next_inside_table and name == "t" and current.text:
            parts.append(current.text)
        elif not next_inside_table and name in {"lineBreak", "break"}:
            parts.append("\n")
        elif not next_inside_table and name == "tab":
            parts.append("\t")

        for child in list(current):
            walk(child, inside_table=next_inside_table)
            if not next_inside_table and child.tail and child.tail.strip():
                parts.append(child.tail)

    walk(node, inside_table=False)
    return _normalize_text("".join(parts))


def _extract_paragraph_rendered_text(node: ET.Element, *, exclude_tables: bool, honor_line_segments: bool = False) -> str:
    text = _extract_paragraph_text_excluding_tables(node) if exclude_tables else _extract_paragraph_text(node)
    # HWPX lineseg/textpos 는 실제 편집기 렌더링 폭, 그림 컨트롤, 자간 계산까지 반영된 결과라
    # 텍스트만 추출한 뒤 그대로 역적용하면 문장 중간이 쉽게 틀어진다.
    # 우선은 명시적 lineBreak/break 와 paragraph 경계만 신뢰하는 보수 경로를 사용한다.
    return text


def _is_decorative_status_text(text: str) -> bool:
    normalized = _normalize_text(text)
    if not normalized:
        return False
    compact = normalized.replace(" ", "")
    return compact in {
        "신청불가",
        "신청가능",
        "(신청불가)",
        "(신청가능)",
        "신청불가신청가능",
        "(신청불가)(신청가능)",
    }


def _extract_inline_box_texts(node: ET.Element) -> list[str]:
    inline_box_texts: list[str] = []
    for child in node.iter():
        if _localname(child.tag) != "tbl":
            continue
        matrix = _extract_table_matrix(child)
        if len(matrix) != 1:
            continue
        if len(matrix[0]) != 1:
            continue
        flattened = _flatten_table_row(matrix[0])
        if not flattened:
            continue
        if len(flattened) > 30:
            continue
        if _is_decorative_status_text(flattened):
            continue
        inline_box_texts.append(flattened)
    return inline_box_texts


def _merge_inline_box_texts(base_text: str, inline_box_texts: list[str], *, left_bracket: str = "(", right_bracket: str = ")") -> str:
    merged = _normalize_text(base_text)
    if not inline_box_texts:
        return merged

    for inline_box_text in inline_box_texts:
        if not inline_box_text:
            continue
        parenthesized = f"{left_bracket}{inline_box_text}{right_bracket}"
        if parenthesized in merged:
            continue
        merged = f"{merged} {parenthesized}".strip() if merged else parenthesized
    return _normalize_text(merged)


def _dedup_multiline(text: str) -> str:
    lines = [line.strip() for line in (text or "").splitlines() if line.strip()]
    deduped: list[str] = []
    for line in lines:
        if not deduped or deduped[-1] != line:
            deduped.append(line)
    return "\n".join(deduped)


def _normalize_table_width(table: list[list[str]]) -> list[list[str]]:
    if not table:
        return table
    max_cols = max(len(row) for row in table)
    return [row + [""] * (max_cols - len(row)) for row in table]


def _is_probable_header_row(row: list[str]) -> bool:
    cells = [_normalize_text(cell) for cell in row]
    non_empty = [cell for cell in cells if cell]
    if len(non_empty) < 2:
        return False

    long_cell_count = sum(1 for cell in non_empty if len(cell) >= 40)
    bullet_like_count = sum(1 for cell in non_empty if cell.startswith(("○", "◦", "-", "*", "※")))
    if long_cell_count >= max(1, len(non_empty) // 2):
        return False
    if bullet_like_count >= max(1, len(non_empty) // 2):
        return False
    return True


def _table_cell_to_markdown_text(cell: str) -> str:
    text = _normalize_text(cell)
    if not text:
        return ""
    return text.replace("|", "\\|").replace("\n", "<br>")


def _table_to_markdown(table: list[list[str]]) -> str:
    table = _normalize_table_width(table)
    if not table:
        return ""

    use_first_row_as_header = _is_probable_header_row(table[0])
    if use_first_row_as_header:
        header_row = table[0]
        body_rows = table[1:]
    else:
        header_row = [f"컬럼{idx + 1}" for idx in range(len(table[0]))]
        body_rows = table

    lines = [
        "| " + " | ".join(_table_cell_to_markdown_text(cell) for cell in header_row) + " |",
        "| " + " | ".join(["---"] * len(header_row)) + " |",
    ]
    for row in body_rows:
        lines.append("| " + " | ".join(_table_cell_to_markdown_text(cell) for cell in row) + " |")
    return "\n".join(lines)


def _is_inside_nested_table_before_cell(
    node: ET.Element,
    cell: ET.Element,
    parent_map: dict[ET.Element, ET.Element],
) -> bool:
    current = parent_map.get(node)
    while current is not None:
        if current is cell:
            return False
        if _localname(current.tag) == "tbl":
            return True
        current = parent_map.get(current)
    return False


def _get_cell_span(node: ET.Element, name: str, default: int = 1) -> int:
    raw = _get_attr(node, name)
    if raw is not None:
        try:
            return max(1, int(str(raw)))
        except (TypeError, ValueError):
            pass

    for child in list(node):
        if _localname(child.tag) != "cellSpan":
            continue
        raw = _get_attr(child, name)
        if raw is None:
            continue
        try:
            return max(1, int(str(raw)))
        except (TypeError, ValueError):
            return default
    return default


def _is_nested_table_inside_cell(
    node: ET.Element,
    parent_map: dict[ET.Element, ET.Element],
) -> bool:
    parent = parent_map.get(node)
    while parent is not None:
        parent_name = _localname(parent.tag)
        if parent_name == "tc":
            return True
        if parent_name == "tbl":
            return True
        parent = parent_map.get(parent)
    return False


def _extract_table_cell_text(node: ET.Element) -> str:
    parent_map = {child: parent for parent in node.iter() for child in parent}
    parts: list[str] = []

    for child in node.iter():
        if child is node:
            continue
        name = _localname(child.tag)

        if name == "p":
            if _is_inside_nested_table_before_cell(child, node, parent_map):
                continue
            paragraph_text = _extract_paragraph_rendered_text(
                child,
                exclude_tables=True,
                honor_line_segments=True,
            )
            if paragraph_text:
                parts.append(paragraph_text)
            continue

        if name == "tbl":
            if _is_inside_nested_table_before_cell(child, node, parent_map):
                continue
            nested_matrix = _extract_table_matrix(child)
            if nested_matrix:
                if len(nested_matrix) == 1:
                    parts.append(_preserve_table_row_text(nested_matrix[0]))
                else:
                    parts.append(_table_to_markdown(nested_matrix))

    return _dedup_multiline("\n".join(part for part in parts if part))


def _normalize_color(value: Any) -> str:
    text = str(value or "").strip().lstrip("#").upper()
    if len(text) == 6 and re.fullmatch(r"[0-9A-F]{6}", text):
        return text
    return ""


def _is_white_like_color(value: Any) -> bool:
    color = _normalize_color(value)
    return color in {"", "FFFFFF", "FFFDF8", "FDFDFD", "FCFCFC", "FAFAFA"}


def _read_header_fill_colors(archive: zipfile.ZipFile) -> dict[str, str]:
    candidates = ["Contents/Header.xml", "Contents/header.xml", "Header.xml", "header.xml"]
    header_xml: bytes | None = None
    for name in candidates:
        try:
            header_xml = archive.read(name)
            break
        except KeyError:
            continue
    if header_xml is None:
        return {}

    try:
        root = ET.fromstring(header_xml)
    except ET.ParseError:
        return {}

    color_map: dict[str, str] = {}
    for node in root.iter():
        if _localname(node.tag) != "borderFill":
            continue
        border_fill_id = str(_get_attr(node, "id") or "").strip()
        if not border_fill_id:
            continue
        face_color = ""
        for child in node.iter():
            if _localname(child.tag) == "winBrush":
                face_color = _normalize_color(_get_attr(child, "faceColor"))
                if face_color:
                    break
        if face_color:
            color_map[border_fill_id] = face_color
    return color_map


def _extract_table_cell_info(node: ET.Element, fill_color_map: dict[str, str]) -> dict[str, Any]:
    border_fill_id = str(_get_attr(node, "borderFillIDRef") or "").strip()
    return {
        "text": _extract_table_cell_text(node),
        "border_fill_id": border_fill_id,
        "fill_color": fill_color_map.get(border_fill_id, ""),
    }


def _empty_table_cell_info() -> dict[str, Any]:
    return {
        "text": "",
        "border_fill_id": "",
        "fill_color": "",
    }


def _is_spacer_column_value(value: str) -> bool:
    normalized = _normalize_text(value)
    if not normalized:
        return True
    return bool(re.fullmatch(r"[↓⇓↑⇑→←↔↕|│┃ ]+", normalized))


def _strip_spacer_columns(matrix: list[list[str]]) -> list[list[str]]:
    if not matrix:
        return matrix
    max_cols = max((len(row) for row in matrix), default=0)
    if max_cols <= 1:
        return matrix

    keep_indices: list[int] = []
    for col_idx in range(max_cols):
        column_values = [row[col_idx] if col_idx < len(row) else "" for row in matrix]
        if all(_is_spacer_column_value(value) for value in column_values):
            continue
        keep_indices.append(col_idx)

    if len(keep_indices) == max_cols or not keep_indices:
        return matrix
    return [
        [row[idx] if idx < len(row) else "" for idx in keep_indices]
        for row in matrix
    ]


def _strip_spacer_columns_from_infos(matrix: list[list[dict[str, Any]]]) -> list[list[dict[str, Any]]]:
    if not matrix:
        return matrix
    max_cols = max((len(row) for row in matrix), default=0)
    if max_cols <= 1:
        return matrix

    keep_indices: list[int] = []
    for col_idx in range(max_cols):
        column_values = [
            (row[col_idx].get("text") if col_idx < len(row) else "")
            for row in matrix
        ]
        if all(_is_spacer_column_value(str(value or "")) for value in column_values):
            continue
        keep_indices.append(col_idx)

    if len(keep_indices) == max_cols or not keep_indices:
        return matrix
    return [
        [row[idx] if idx < len(row) else {"text": "", "border_fill_id": "", "fill_color": ""} for idx in keep_indices]
        for row in matrix
    ]


def _extract_table_cell_info_matrix(node: ET.Element, fill_color_map: dict[str, str]) -> list[list[dict[str, Any]]]:
    matrix: list[list[dict[str, Any]]] = []
    row_spans: dict[int, dict[str, Any]] = {}
    rows = [child for child in list(node) if _localname(child.tag) == "tr"]

    for row_node in rows:
        row_cells: list[dict[str, Any]] = []
        col_idx = 0
        table_cells = [child for child in list(row_node) if _localname(child.tag) == "tc"]
        tc_idx = 0

        while tc_idx < len(table_cells) or col_idx in row_spans:
            if col_idx in row_spans:
                span_info = row_spans[col_idx]
                row_cells.append(dict(span_info["cell"]))
                span_info["remaining"] -= 1
                if span_info["remaining"] <= 0:
                    del row_spans[col_idx]
                col_idx += 1
                continue

            if tc_idx >= len(table_cells):
                break

            cell_node = table_cells[tc_idx]
            tc_idx += 1

            cell_info = _extract_table_cell_info(cell_node, fill_color_map)
            col_span = _get_cell_span(cell_node, "colSpan", 1)
            row_span = _get_cell_span(cell_node, "rowSpan", 1)

            for offset in range(col_span):
                current_col = col_idx + offset
                value = dict(cell_info) if offset == 0 else {
                    "text": "",
                    "border_fill_id": str(cell_info.get("border_fill_id") or ""),
                    "fill_color": str(cell_info.get("fill_color") or ""),
                }
                row_cells.append(value)

                if row_span > 1:
                    repeated = dict(cell_info)
                    if offset > 0:
                        repeated["text"] = ""
                    row_spans[current_col] = {
                        "cell": repeated,
                        "remaining": row_span - 1,
                    }

            col_idx += col_span

        if any(str(cell.get("text") or "").strip() for cell in row_cells):
            matrix.append(row_cells)
    return _strip_spacer_columns_from_infos(matrix)


def _extract_table_matrix(node: ET.Element) -> list[list[str]]:
    matrix = [
        [str(cell.get("text") or "") for cell in row]
        for row in _extract_table_cell_info_matrix(node, {})
    ]
    return _strip_spacer_columns(matrix)


def _extract_table_rows(node: ET.Element) -> list[str]:
    return [" | ".join(cell for cell in row if cell.strip()) for row in _extract_table_matrix(node)]


def _flatten_table_row(row: list[str]) -> str:
    return _normalize_text(" ".join(cell for cell in row if _normalize_text(cell)))


def _preserve_table_row_text(row: list[str]) -> str:
    normalized_cells = [_normalize_text(cell) for cell in row if _normalize_text(cell)]
    if not normalized_cells:
        return ""
    if len(normalized_cells) == 1:
        return normalized_cells[0]
    return "\n".join(normalized_cells)


def _is_caption_like_text(text: str) -> bool:
    normalized = _normalize_text(text)
    if not normalized:
        return False
    return bool(re.fullmatch(r"[《<〈「『].+[》>〉」』]", normalized))


def _extract_captioned_table_content(matrix: list[list[str]]) -> tuple[str, str] | None:
    if len(matrix) < 2:
        return None

    first_row_cells = [_normalize_text(cell) for cell in matrix[0] if _normalize_text(cell)]
    if len(first_row_cells) != 1:
        return None

    caption = first_row_cells[0]
    if not _is_caption_like_text(caption):
        return None

    body_lines: list[str] = []
    for row in matrix[1:]:
        row_text = _preserve_table_row_text(row)
        if row_text:
            body_lines.append(row_text)

    if not body_lines:
        return None
    if body_lines and _normalize_text(body_lines[0]) == caption:
        body_lines = body_lines[1:]
    if not body_lines:
        return caption, caption
    return caption, "\n".join([caption, *body_lines]).strip()


def _has_substantive_text(value: str) -> bool:
    normalized = _normalize_text(value)
    if not normalized:
        return False
    if normalized in {"-", "–", "—", "~"}:
        return True
    return any(char.isalnum() for char in normalized)


def _row_has_substantive_cells(row: list[str]) -> bool:
    return any(_has_substantive_text(cell) for cell in row)


def _normalize_header_cell_text(value: str) -> str:
    return " ".join(
        _normalize_text(line)
        for line in _normalize_text(value).splitlines()
        if _normalize_text(line)
    ).strip()


def _format_grouped_header_row_content(headers: list[str], row: list[str]) -> str:
    """
    병합된 헤더행은 빈 헤더 셀이 앞선 헤더의 continuation 인 경우가 많다.
    예: ["구분", "", "논물관리", "바이오차", "가을갈이"]
    이때 첫 두 컬럼 값은 `구분: A법인·단체 | a농업법인` 형태로 묶는다.
    """

    parts: list[str] = []
    idx = 0
    row_len = max(len(headers), len(row))

    while idx < row_len:
        header = _normalize_header_cell_text(headers[idx]) if idx < len(headers) else ""
        if not header:
            idx += 1
            continue

        group_values: list[str] = []
        group_idx = idx
        while group_idx < row_len:
            if group_idx > idx:
                next_header = _normalize_header_cell_text(headers[group_idx]) if group_idx < len(headers) else ""
                if next_header:
                    break

            cell_value = _normalize_text(row[group_idx]) if group_idx < len(row) else ""
            if cell_value and _has_substantive_text(cell_value):
                group_values.append(cell_value)
            group_idx += 1

        if group_values:
            parts.append(f"{header}: {' | '.join(group_values)}")
        idx = group_idx

    return "\n".join(parts).strip()


def _format_table_row_content(headers: list[str], row: list[str]) -> str:
    if any(not _normalize_header_cell_text(header) for header in headers):
        grouped = _format_grouped_header_row_content(headers, row)
        if grouped:
            return grouped

    parts: list[str] = []
    for idx, value in enumerate(row):
        cell_value = _normalize_text(value)
        if not cell_value:
            continue
        if not _has_substantive_text(cell_value):
            continue
        value_lines = [_normalize_text(line) for line in cell_value.splitlines() if _normalize_text(line)]
        if not value_lines:
            continue
        header = _normalize_header_cell_text(headers[idx]) if idx < len(headers) else ""
        if header:
            first_line = value_lines[0]
            remaining_lines = value_lines[1:]
            merged = f"{header}: {first_line}"
            if remaining_lines:
                merged = "\n".join([merged, *remaining_lines])
            parts.append(merged)
        else:
            parts.append("\n".join(value_lines))
    return "\n".join(parts).strip()


def _looks_like_paired_header_value_table(matrix: list[list[str]]) -> bool:
    if len(matrix) < 2:
        return False
    max_cols = max((len(row) for row in matrix), default=0)
    if max_cols < 2 or max_cols % 2 != 0:
        return False

    meaningful_pairs = 0
    for row in matrix:
        row_len = len(row)
        for idx in range(0, row_len - 1, 2):
            header = _normalize_text(row[idx])
            value = _normalize_text(row[idx + 1])
            if header and value:
                meaningful_pairs += 1
    return meaningful_pairs >= max(2, len(matrix))


def _looks_like_header_value_pair_group_row(row: list[str]) -> bool:
    """
    패턴 2-1 추정:
    헤더 - (소헤더:값 목록) 형태는 값 셀 안에 여러 줄 목록이나 내부 콜론 구조가 남는 경우가 많다.
    """

    meaningful_cells = [_normalize_text(cell) for cell in row if _normalize_text(cell)]
    if len(meaningful_cells) >= 3 and len(meaningful_cells) % 2 == 1:
        return True

    row_len = len(row)
    for idx in range(0, row_len - 1, 2):
        value = _normalize_text(row[idx + 1])
        if not value:
            continue
        value_lines = [_normalize_text(line) for line in value.splitlines() if _normalize_text(line)]
        if len(value_lines) >= 2:
            return True
        if any(token in value for token in (":", "：")):
            return True
    return False


def _infer_header_value_pattern_from_row(row: list[str]) -> TablePattern:
    if _looks_like_header_value_pair_group_row(row):
        return TablePattern.HEADER_VALUE_PAIR_GROUPS
    return TablePattern.HEADER_VALUE_PAIRS


def _looks_like_paired_header_value_table_by_color(cell_info_matrix: list[list[dict[str, Any]]]) -> bool:
    if len(cell_info_matrix) < 2:
        return False
    max_cols = max((len(row) for row in cell_info_matrix), default=0)
    if max_cols < 2 or max_cols % 2 != 0:
        return False

    colored_header_cells = 0
    plain_value_cells = 0
    mismatches = 0

    for row in cell_info_matrix:
        for idx in range(0, min(len(row) - 1, max_cols - 1), 2):
            header_info = row[idx]
            value_info = row[idx + 1]
            header_text = _normalize_text(str(header_info.get("text") or ""))
            value_text = _normalize_text(str(value_info.get("text") or ""))
            if not header_text or not value_text:
                continue
            if not _is_white_like_color(header_info.get("fill_color")):
                colored_header_cells += 1
            else:
                mismatches += 1
            if _is_white_like_color(value_info.get("fill_color")):
                plain_value_cells += 1
            else:
                mismatches += 1

    return colored_header_cells >= 2 and plain_value_cells >= 2 and mismatches <= max(2, colored_header_cells // 2)


def _first_row_all_colored(cell_info_matrix: list[list[dict[str, Any]]]) -> bool:
    if not cell_info_matrix or not cell_info_matrix[0]:
        return False
    first_row = cell_info_matrix[0]
    visible_cells = [cell for cell in first_row if _normalize_text(str(cell.get("text") or ""))]
    if len(visible_cells) < 2:
        return False
    return all(not _is_white_like_color(cell.get("fill_color")) for cell in visible_cells)


def _first_row_suggests_paired_layout(cell_info_matrix: list[list[dict[str, Any]]]) -> bool:
    if len(cell_info_matrix) < 2:
        return False
    max_cols = max((len(row) for row in cell_info_matrix), default=0)
    if max_cols < 2 or max_cols % 2 != 0:
        return False
    if _first_row_all_colored(cell_info_matrix):
        return False

    score = 0
    first_row = cell_info_matrix[0]
    for idx in range(0, min(len(first_row) - 1, max_cols - 1), 2):
        header_info = first_row[idx]
        value_info = first_row[idx + 1]
        header_text = _normalize_text(str(header_info.get("text") or ""))
        value_text = _normalize_text(str(value_info.get("text") or ""))
        if not header_text or not value_text:
            continue
        if not _is_white_like_color(header_info.get("fill_color")):
            score += 2
        if _is_white_like_color(value_info.get("fill_color")):
            score += 1
    return score >= 2


def _row_visible_color_flags(row_infos: list[dict[str, Any]]) -> list[bool]:
    flags: list[bool] = []
    for cell in row_infos:
        if not _normalize_text(str(cell.get("text") or "")):
            continue
        flags.append(not _is_white_like_color(cell.get("fill_color")))
    return flags


def _is_first_visible_cell_colored(row_infos: list[dict[str, Any]]) -> bool:
    for cell in row_infos:
        if not _normalize_text(str(cell.get("text") or "")):
            continue
        return not _is_white_like_color(cell.get("fill_color"))
    return False


def _infer_row_color_driven_pattern(row_infos: list[dict[str, Any]]) -> TablePattern | None:
    """
    행 단위 색 패턴 판정 규칙:
    - 1번째 의미 있는 컬럼: 유색이어야 함
    - 2번째 의미 있는 컬럼:
      - 흰색 => 일반 헤더-값 pair
      - 유색 => 3번째 의미 있는 컬럼 확인
    - 3번째 의미 있는 컬럼:
      - 흰색 => 2-1. 헤더 : [소헤더-값, ...]
      - 유색 => 컬럼 헤더 행으로 전환
    """

    color_flags = _row_visible_color_flags(row_infos)
    if len(color_flags) < 2:
        return None
    if not color_flags[0]:
        return None
    if not color_flags[1]:
        return TablePattern.HEADER_VALUE_PAIRS
    if len(color_flags) < 3:
        return None
    if not color_flags[2]:
        return TablePattern.HEADER_VALUE_PAIR_GROUPS
    return None


def _row_paired_layout_score(row_infos: list[dict[str, Any]]) -> int:
    # 헤더-값 계열 표는 첫 번째 의미 있는 셀이 반드시 유색 헤더여야 한다.
    # 이 조건이 없으면 일반 헤더행 표의 데이터행도 일부 칸만 보고 pair 로 오인하기 쉽다.
    visible_cells = [cell for cell in row_infos if _normalize_text(str(cell.get("text") or ""))]
    if len(visible_cells) < 2:
        return 0

    first_cell = visible_cells[0]
    if _is_white_like_color(first_cell.get("fill_color")):
        return 0

    score = 0
    row_len = len(row_infos)
    for idx in range(0, row_len - 1, 2):
        header_info = row_infos[idx]
        value_info = row_infos[idx + 1]
        header_text = _normalize_text(str(header_info.get("text") or ""))
        value_text = _normalize_text(str(value_info.get("text") or ""))
        if not header_text or not value_text:
            continue
        # 헤더 위치 셀은 유색, 값 위치 셀은 비유색일 때만 페어 점수를 준다.
        if _is_white_like_color(header_info.get("fill_color")):
            continue
        score += 2
        if _is_white_like_color(value_info.get("fill_color")):
            score += 1
    return score


def _classify_table_row_layout(
    row_infos: list[dict[str, Any]],
    *,
    expected_pattern: TablePattern,
    is_first_row: bool = False,
) -> TableRowRole:
    visible_cells = [cell for cell in row_infos if _normalize_text(str(cell.get("text") or ""))]
    if not visible_cells:
        return TableRowRole.EMPTY

    has_colored = any(not _is_white_like_color(cell.get("fill_color")) for cell in visible_cells)
    all_colored = has_colored and all(not _is_white_like_color(cell.get("fill_color")) for cell in visible_cells)
    paired_score = _row_paired_layout_score(row_infos)
    inferred_color_pattern = _infer_row_color_driven_pattern(row_infos)

    if is_first_row:
        if not has_colored:
            return TableRowRole.HEADER_ROW
        if all_colored:
            return TableRowRole.HEADER_ROW
        if inferred_color_pattern in {
            TablePattern.HEADER_VALUE_PAIRS,
            TablePattern.HEADER_VALUE_PAIR_GROUPS,
        }:
            return TableRowRole.HEADER_VALUE_ROW
        if paired_score >= 2:
            return TableRowRole.HEADER_VALUE_ROW
        return TableRowRole.HEADER_ROW if expected_pattern == TablePattern.HEADER_ROW_RECORDS else TableRowRole.DATA_ROW

    if all_colored:
        return TableRowRole.HEADER_ROW
    if expected_pattern == TablePattern.HEADER_ROW_RECORDS:
        if inferred_color_pattern == TablePattern.HEADER_VALUE_PAIRS:
            return TableRowRole.HEADER_VALUE_ROW
        return TableRowRole.DATA_ROW
    if expected_pattern in {
        TablePattern.HEADER_VALUE_PAIRS,
        TablePattern.HEADER_VALUE_PAIR_GROUPS,
    }:
        if inferred_color_pattern in {
            TablePattern.HEADER_VALUE_PAIRS,
            TablePattern.HEADER_VALUE_PAIR_GROUPS,
        }:
            return TableRowRole.HEADER_VALUE_ROW
        if inferred_color_pattern is None and has_colored:
            return TableRowRole.HEADER_ROW
    if paired_score >= 2:
        return TableRowRole.HEADER_VALUE_ROW
    return TableRowRole.DATA_ROW


def _format_paired_header_value_row(row: list[str]) -> str:
    parts: list[str] = []
    row_len = len(row)
    for idx in range(0, row_len - 1, 2):
        header = _normalize_header_cell_text(row[idx])
        value = _normalize_text(row[idx + 1])
        if not header or not value:
            continue
        value_lines = [_normalize_text(line) for line in value.splitlines() if _normalize_text(line)]
        if not value_lines:
            continue
        merged = f"{header}: {value_lines[0]}"
        if len(value_lines) > 1:
            merged = "\n".join([merged, *value_lines[1:]])
        parts.append(merged)
    return "\n".join(parts).strip()


def _format_leading_header_value_row(row: list[str]) -> str:
    meaningful_cells = [_normalize_text(cell) for cell in row if _normalize_text(cell)]
    if len(meaningful_cells) < 2:
        return ""
    header = _normalize_header_cell_text(meaningful_cells[0])
    value = ", ".join(meaningful_cells[1:]).strip()
    if not header or not value:
        return ""
    return f"{header}: {value}"


def _format_header_value_pair_group_row(row: list[str]) -> str:
    normalized_cells = [_normalize_text(cell) for cell in row if _normalize_text(cell)]
    if len(normalized_cells) < 3:
        return _format_paired_header_value_row(row)

    main_header = _normalize_header_cell_text(normalized_cells[0])
    sub_parts: list[str] = []
    idx = 1
    while idx < len(normalized_cells):
        sub_header = _normalize_header_cell_text(normalized_cells[idx])
        sub_value = normalized_cells[idx + 1] if idx + 1 < len(normalized_cells) else ""
        if sub_header and sub_value:
            sub_parts.append(f"{sub_header} : {sub_value}")
        elif sub_header:
            sub_parts.append(sub_header)
        idx += 2

    if not main_header:
        return ", ".join(sub_parts).strip()
    if not sub_parts:
        return main_header
    return f"{main_header}: {', '.join(sub_parts)}".strip()


def _count_table_cells(node: ET.Element) -> int:
    count = 0
    for child in node.iter():
        if _localname(child.tag) == "tc":
            count += 1
    return count


def _is_short_inline_box_table(node: ET.Element, parent_map: dict[ET.Element, ET.Element]) -> bool:
    parent = parent_map.get(node)
    inside_paragraph = False
    while parent is not None:
        if _localname(parent.tag) == "p":
            inside_paragraph = True
            break
        parent = parent_map.get(parent)
    if not inside_paragraph:
        return False

    matrix = _extract_table_matrix(node)
    if len(matrix) != 1 or len(matrix[0]) != 1:
        return False
    flattened = _flatten_table_row(matrix[0])
    return bool(flattened) and len(flattened) <= 30


def _extract_hwpx_events(content: bytes) -> list[dict[str, Any]]:
    try:
        archive = zipfile.ZipFile(io.BytesIO(content))
    except zipfile.BadZipFile as exc:
        raise RagEmbeddingParseError("유효한 HWPX zip 구조가 아닙니다.") from exc

    fill_color_map = _read_header_fill_colors(archive)

    section_files = sorted(
        [name for name in archive.namelist() if re.match(r"^Contents/section\d+\.xml$", name)],
        key=lambda path: int(re.search(r"section(\d+)\.xml$", path).group(1)) if re.search(r"section(\d+)\.xml$", path) else 999999,
    )
    if not section_files:
        raise RagEmbeddingParseError("HWPX section XML을 찾지 못했습니다.")

    events: list[dict[str, Any]] = []
    source_order = 0
    for section_name in section_files:
        xml_bytes = archive.read(section_name)
        parser = ET.XMLParser()
        try:
            root = ET.fromstring(xml_bytes, parser=parser)
        except ET.ParseError:
            continue
        parent_map = {child: parent for parent in root.iter() for child in parent}
        for node in root.iter():
            name = _localname(node.tag)
            if name == "p":
                parent = parent_map.get(node)
                location = "paragraph"
                enclosing_table: ET.Element | None = None
                while parent is not None:
                    parent_name = _localname(parent.tag)
                    if parent_name == "tbl":
                        location = "table-cell"
                        enclosing_table = parent
                        break
                    parent = parent_map.get(parent)
                text = _extract_paragraph_rendered_text(
                    node,
                    exclude_tables=True,
                    honor_line_segments=False,
                )
                inline_box_texts = _extract_inline_box_texts(node)
                text = _merge_inline_box_texts(text, inline_box_texts)
                if not text:
                    continue
                if _is_decorative_status_text(text):
                    continue
                # 표는 별도 row/table 이벤트로 대표시킨다.
                if location == "table-cell" and enclosing_table is not None:
                    continue
                source_order += 1
                events.append(
                    {
                        "source_order": source_order,
                        "text": text,
                        "location": location,
                        "block_type": "paragraph" if location == "paragraph" else "table-cell-paragraph",
                        "section": section_name,
                    }
                )
            elif name == "tbl":
                if _is_nested_table_inside_cell(node, parent_map):
                    continue
                if _is_short_inline_box_table(node, parent_map):
                    continue
                cell_info_matrix = _extract_table_cell_info_matrix(node, fill_color_map)
                matrix = [
                    [str(cell.get("text") or "") for cell in row]
                    for row in cell_info_matrix
                ]
                if not matrix:
                    continue
                cell_count = _count_table_cells(node)
                if len(matrix) == 1:
                    flattened = _flatten_table_row(matrix[0])
                    preserved = _preserve_table_row_text(matrix[0])
                    if flattened and preserved:
                        source_order += 1
                        events.append(
                            {
                                "source_order": source_order,
                                "text": preserved,
                                "match_text": flattened,
                                "location": "table",
                                "block_type": "table-single-row",
                                "section": section_name,
                                "metadata": {
                                    "row_count": 1,
                                    "cell_count": cell_count,
                                },
                            }
                        )
                    continue
                if cell_count == 1:
                    source_order += 1
                    events.append(
                        {
                            "source_order": source_order,
                            "text": "\n".join(" | ".join(cell for cell in row if cell.strip()) for row in matrix),
                            "location": "table",
                            "block_type": "table",
                            "section": section_name,
                            "metadata": {
                                "row_count": len(matrix),
                                "cell_count": cell_count,
                            },
                        }
                    )
                    continue

                captioned_table = _extract_captioned_table_content(matrix)
                if captioned_table is not None:
                    caption_text, full_text = captioned_table
                    source_order += 1
                    events.append(
                        {
                            "source_order": source_order,
                            "text": full_text,
                            "location": "table",
                            "block_type": "table",
                            "section": section_name,
                            "metadata": {
                                "row_count": len(matrix),
                                "cell_count": cell_count,
                                "table_caption": caption_text,
                                "table_caption_position": "before-table",
                                "table_caption_embedded": True,
                            },
                        }
                    )
                    continue

                # 표 패턴 분류 기준
                # 1. HEADER_ROW_RECORDS: 첫째 행 헤더, 나머지 행 데이터
                # 2. HEADER_VALUE_PAIRS: 헤더-값 쌍 반복
                # 3. HEADER_VALUE_PAIR_GROUPS: 헤더-(소헤더:값 목록) 변형
                default_pattern = TablePattern.HEADER_ROW_RECORDS
                if (
                    _first_row_suggests_paired_layout(cell_info_matrix)
                    or _looks_like_paired_header_value_table_by_color(cell_info_matrix)
                    or _looks_like_paired_header_value_table(matrix)
                ):
                    default_pattern = _infer_header_value_pattern_from_row(matrix[0])

                current_pattern = default_pattern
                current_headers: list[str] = []
                emitted = False
                force_record_rows_after_header = False

                for row_index, (row_infos, row) in enumerate(zip(cell_info_matrix, matrix), start=1):
                    if not _row_has_substantive_cells(row):
                        continue

                    color_pattern = _infer_row_color_driven_pattern(row_infos)
                    row_role = _classify_table_row_layout(
                        row_infos,
                        expected_pattern=current_pattern,
                        is_first_row=(row_index == 1),
                    )

                    if current_pattern == TablePattern.HEADER_ROW_RECORDS and current_headers:
                        if row_index == 2:
                            row_role = TableRowRole.DATA_ROW
                            if _is_first_visible_cell_colored(row_infos):
                                force_record_rows_after_header = True
                        elif force_record_rows_after_header:
                            row_role = TableRowRole.DATA_ROW

                    if row_role == TableRowRole.HEADER_ROW:
                        current_headers = [_normalize_header_cell_text(cell) for cell in row]
                        current_pattern = TablePattern.HEADER_ROW_RECORDS
                        force_record_rows_after_header = False
                        continue

                    if row_role == TableRowRole.HEADER_VALUE_ROW:
                        active_pair_pattern = color_pattern or _infer_header_value_pattern_from_row(row)
                        if (
                            current_pattern == TablePattern.HEADER_ROW_RECORDS
                            and color_pattern == TablePattern.HEADER_VALUE_PAIRS
                        ):
                            row_text = _format_leading_header_value_row(row)
                        elif active_pair_pattern == TablePattern.HEADER_VALUE_PAIR_GROUPS:
                            row_text = _format_header_value_pair_group_row(row)
                        else:
                            row_text = _format_paired_header_value_row(row)
                        if not row_text:
                            continue
                        current_pattern = active_pair_pattern
                        source_order += 1
                        emitted = True
                        events.append(
                            {
                                "source_order": source_order,
                                "text": row_text,
                                "location": "table",
                                "block_type": "table-row",
                                "section": section_name,
                                "metadata": {
                                    "row_count": len(matrix),
                                    "cell_count": cell_count,
                                    "row_index": row_index,
                                    "table_pattern": current_pattern.value,
                                    "table_row_role": row_role.value,
                                },
                            }
                        )
                        continue

                    row_text = ""
                    row_pattern = current_pattern
                    if current_headers:
                        row_text = _format_table_row_content(current_headers, row)
                        row_pattern = TablePattern.HEADER_ROW_RECORDS
                    elif current_pattern in {
                        TablePattern.HEADER_VALUE_PAIRS,
                        TablePattern.HEADER_VALUE_PAIR_GROUPS,
                    }:
                        row_pattern = color_pattern or _infer_header_value_pattern_from_row(row)
                        if row_pattern == TablePattern.HEADER_VALUE_PAIR_GROUPS:
                            row_text = _format_header_value_pair_group_row(row)
                        else:
                            row_text = _format_paired_header_value_row(row)
                    if not row_text:
                        row_text = "\n".join(cell for cell in row if _normalize_text(cell))
                    if not row_text:
                        continue

                    source_order += 1
                    emitted = True
                    events.append(
                        {
                            "source_order": source_order,
                            "text": row_text,
                            "location": "table",
                            "block_type": "table-row",
                            "section": section_name,
                            "metadata": {
                                "row_count": len(matrix),
                                "cell_count": cell_count,
                                "row_index": row_index,
                                "header_row": current_headers if current_headers else None,
                                "table_pattern": row_pattern.value,
                                "table_row_role": TableRowRole.DATA_ROW.value,
                            },
                        }
                    )

                if emitted:
                    continue

                source_order += 1
                events.append(
                    {
                        "source_order": source_order,
                        "text": "\n".join(" | ".join(cell for cell in row if cell.strip()) for row in matrix),
                        "location": "table",
                        "block_type": "table",
                        "section": section_name,
                        "metadata": {
                            "row_count": len(matrix),
                            "cell_count": cell_count,
                        },
                    }
                )
    return events


def _compile_heading_levels(schema: dict[str, Any]) -> list[dict[str, Any]]:
    normalized = _normalize_schema(schema)
    levels = normalized.get("levels")
    if not isinstance(levels, list):
        return []
    compiled: list[dict[str, Any]] = []
    for idx, level in enumerate(levels, start=1):
        if not isinstance(level, dict):
            continue
        pattern = _compile_level_pattern(level)
        rule_type = str(level.get("rule_type") or "").strip()
        if not pattern and rule_type != "appendix_title_table":
            continue
        regex: re.Pattern[str] | None = None
        if pattern:
            try:
                regex = re.compile(pattern)
            except re.error:
                continue
        compiled.append(
            {
                "depth": int(level.get("depth") or idx),
                "location": _normalize_level_location(level.get("location") or "paragraph"),
                "rule_id": str(level.get("rule_id") or "").strip(),
                "rule_type": rule_type,
                "notation": str(level.get("notation") or "").strip(),
                "pattern": regex,
                "rule_options": _get_rule_options(level),
            }
        )
    compiled.sort(key=lambda item: int(item["depth"]))
    return compiled


def _matches_appendix_title_table_runtime(text: str, location: str, level: dict[str, Any], metadata: dict[str, Any] | None) -> bool:
    if location != "table":
        return False
    options = _get_rule_options(level)
    normalized_text = _normalize_text(text)
    if not normalized_text:
        return False
    row_count = metadata.get("row_count") if isinstance(metadata, dict) else None
    cell_count = metadata.get("cell_count") if isinstance(metadata, dict) else None
    if _bool_option(options, "single_row_table_only", True) and row_count not in {None, 1}:
        return False
    required_cell_count = options.get("require_cell_count")
    try:
        expected_cell_count = int(required_cell_count) if required_cell_count is not None else None
    except (TypeError, ValueError):
        expected_cell_count = None
    if expected_cell_count is not None and cell_count not in {None, expected_cell_count}:
        return False
    keywords = options.get("keywords")
    keyword_group = (
        "|".join(re.escape(str(keyword).strip()) for keyword in keywords if str(keyword).strip())
        if isinstance(keywords, list)
        else ""
    ) or r"참고|첨부"
    title_min_length = _int_option(options, "title_cell_min_length", 2)
    pattern = (
        rf"^\s*(?:{keyword_group})\s*\d+\s*(?:\||\s)\s*(.+)$"
        if _bool_option(options, "left_cell_number_required", True)
        else rf"^\s*(?:{keyword_group})\s*(?:\||\s)\s*(.+)$"
    )
    match = re.match(pattern, normalized_text)
    if not match:
        return False
    title_text = _normalize_text(match.group(1) or "")
    return len(title_text) >= title_min_length


def _level_matches_text(
    line: str,
    location: str,
    level: dict[str, Any],
    *,
    metadata: dict[str, Any] | None = None,
) -> bool:
    level_location = str(level.get("location") or "paragraph")
    if not _location_matches(level_location, location):
        return False
    rule_type = str(level.get("rule_type") or "").strip()
    if rule_type == "appendix_title_table":
        return _matches_appendix_title_table_runtime(line, location, level, metadata)
    pattern = level.get("pattern")
    return isinstance(pattern, re.Pattern) and pattern.match(line) is not None


def _match_heading(line: str, location: str, compiled_levels: list[dict[str, Any]], *, metadata: dict[str, Any] | None = None) -> int | None:
    for level in compiled_levels:
        if _level_matches_text(line, location, level, metadata=metadata):
            return int(level["depth"])
    return None


def _match_heading_level(
    line: str,
    location: str,
    compiled_levels: list[dict[str, Any]],
    *,
    metadata: dict[str, Any] | None = None,
) -> dict[str, Any] | None:
    for level in compiled_levels:
        if _level_matches_text(line, location, level, metadata=metadata):
            return level
    return None


def _is_symbolic_reentry_level(level: dict[str, Any]) -> bool:
    rule_type = str(level.get("rule_type") or "").strip().lower()
    rule_id = str(level.get("rule_id") or "").strip().lower()
    notation = str(level.get("notation") or "").strip()
    if rule_type in {"symbol", "custom:symbol"} or rule_id in {"symbol", "custom:symbol"}:
        return True
    compact = re.sub(r"[\s,.\-_/|()]+", "", notation)
    return bool(compact and not re.search(r"[0-9A-Za-z가-힣ⅠⅡⅢⅣⅤⅥⅦⅧⅨⅩ①-⑳㉮-㉻]", compact))


def _criteria_targets_level(criteria: dict[str, Any], level: dict[str, Any]) -> bool:
    match = criteria.get("match")
    if not isinstance(match, dict):
        return False

    target_depth = match.get("depth")
    if target_depth is not None:
        try:
            if int(level.get("depth") or 0) > int(target_depth):
                return False
        except (TypeError, ValueError):
            return False

    target_rule_id = str(match.get("rule_id") or "").strip()
    if target_rule_id and str(level.get("rule_id") or "").strip() != target_rule_id:
        return False

    target_notation = str(match.get("notation") or "").strip()
    if target_notation and str(level.get("notation") or "").strip() != target_notation:
        return False

    return True


def _matches_exit_criteria_runtime(
    text: str,
    location: str,
    *,
    criteria: dict[str, Any] | None,
    compiled_levels: list[dict[str, Any]],
    metadata: dict[str, Any] | None = None,
    exclude_symbolic_reentry_levels: bool = False,
) -> bool:
    if not isinstance(criteria, dict):
        return False
    if str(criteria.get("mode") or "").strip() != "matched_heading":
        return False

    for level in compiled_levels:
        if exclude_symbolic_reentry_levels and _is_symbolic_reentry_level(level):
            continue
        if not _criteria_targets_level(criteria, level):
            continue
        if _level_matches_text(text, location, level, metadata=metadata):
            return True
    return False


def _matches_default_appendix_exit_runtime(
    text: str,
    location: str,
    compiled_main_levels: list[dict[str, Any]],
    *,
    metadata: dict[str, Any] | None = None,
) -> bool:
    for level in compiled_main_levels:
        if int(level.get("depth") or 0) > 2:
            continue
        if _is_symbolic_reentry_level(level):
            continue
        if _level_matches_text(text, location, level, metadata=metadata):
            return True
    return False


def _is_main_reentry_level(level: dict[str, Any] | None) -> bool:
    if not isinstance(level, dict):
        return False
    try:
        depth = int(level.get("depth") or 0)
    except (TypeError, ValueError):
        return False
    if depth < 1 or depth > 2:
        return False
    return not _is_symbolic_reentry_level(level)

def _heading_path_item_text(item: dict[str, Any], *, has_descendant: bool = False) -> str:
    text = _normalize_text(str(item.get("text") or ""))
    if not text:
        return ""

    kind = str(item.get("kind") or "title")
    if kind == "title":
        return text

    notation = _normalize_text(str(item.get("notation") or ""))
    if has_descendant and notation:
        return notation

    rule_type = str(item.get("rule_type") or "")
    remainder = _strip_heading_notation(text, rule_type)
    if remainder:
        return text

    return notation or text


def _build_heading_path(stack: list[dict[str, Any]]) -> str:
    parts: list[str] = []
    for idx, item in enumerate(stack):
        current_depth = int(item.get("depth") or 0)
        has_descendant = any(int(next_item.get("depth") or 0) > current_depth for next_item in stack[idx + 1:])
        part = _heading_path_item_text(item, has_descendant=has_descendant)
        if part:
            parts.append(part)
    return " > ".join(parts)


def _build_heading_path_for_content(stack: list[dict[str, Any]], content: str) -> str:
    normalized_content = _normalize_text(content)
    parts: list[str] = []
    for idx, item in enumerate(stack):
        current_depth = int(item.get("depth") or 0)
        has_descendant = any(int(next_item.get("depth") or 0) > current_depth for next_item in stack[idx + 1:])
        text = _normalize_text(str(item.get("text") or ""))
        notation = _normalize_text(str(item.get("notation") or ""))
        kind = str(item.get("kind") or "title")

        if kind == "divider" and text and notation and text in normalized_content:
            part = notation
        else:
            part = _heading_path_item_text(item, has_descendant=has_descendant)

        if part:
            parts.append(part)
    return " > ".join(parts)


def _classify_heading_kind(rule_type: str) -> str:
    normalized = (rule_type or "").strip()
    if normalized in {"symbol", "circled_number", "circled_korean", "paren_numeric", "paren_korean"}:
        return "divider"
    return "title"


def _extract_heading_notation(text: str, rule_type: str) -> str:
    stripped = _normalize_text(text)
    patterns = {
        "roman": r"^([ⅠⅡⅢⅣⅤⅥⅦⅧⅨⅩ](?:\.)?)",
        "numeric_dot": r"^([0-9]+(?:\.)?)",
        "korean_letter_dot": r"^([가-힣](?:\.)?)",
        "numeric_paren": r"^([0-9]+\))",
        "korean_letter_paren": r"^([가-힣]\))",
        "circled_number": r"^([①-⑳])",
        "circled_korean": r"^([㉮-㉻])",
        "paren_numeric": r"^(\([0-9]+\))",
        "paren_korean": r"^(\([가-힣]\))",
        "symbol": r"^([○●□■◇◆❍⊙▪▫◦•·])",
        "legal_article": r"^(제[0-9]+조(?:의[0-9]+)?)",
    }
    pattern = patterns.get((rule_type or "").strip(), "")
    if pattern:
        match = re.match(pattern, stripped)
        if match:
            return _normalize_text(match.group(1))
    return stripped


def _strip_heading_notation(text: str, rule_type: str) -> str:
    stripped = _normalize_text(text)
    patterns = {
        "roman": r"^[ⅠⅡⅢⅣⅤⅥⅦⅧⅨⅩ](?:\.)?\s*",
        "numeric_dot": r"^[0-9]+(?:\.)?\s*",
        "korean_letter_dot": r"^[가-힣](?:\.)?\s*",
        "numeric_paren": r"^[0-9]+\)\s*",
        "korean_letter_paren": r"^[가-힣]\)\s*",
        "circled_number": r"^[①-⑳]\s*",
        "circled_korean": r"^[㉮-㉻]\s*",
        "paren_numeric": r"^\([0-9]+\)\s*",
        "paren_korean": r"^\([가-힣]\)\s*",
        "symbol": r"^[○●□■◇◆❍⊙▪▫◦•·]\s*",
        "legal_article": r"^제[0-9]+조(?:의[0-9]+)?(?:\s*\([^\)]+\))?\s*",
    }
    pattern = patterns.get((rule_type or "").strip(), "")
    if not pattern:
        return stripped
    return _normalize_text(re.sub(pattern, "", stripped, count=1))


def _update_heading_stack(stack: list[dict[str, Any]], depth: int, text: str, *, rule_type: str) -> list[dict[str, Any]]:
    next_stack = [item for item in stack if int(item.get("depth") or 0) < depth]
    next_stack.append(
        {
            "depth": depth,
            "text": text,
            "kind": _classify_heading_kind(rule_type),
            "rule_type": rule_type,
            "notation": _extract_heading_notation(text, rule_type),
        }
    )
    return next_stack


def _heading_has_inline_payload(line: str) -> bool:
    text = _normalize_text(line)
    if not text:
        return False
    if ":" in text or "：" in text:
        return True
    # 제목 한 줄 안에 날짜/기간/금액 같은 실질값이 같이 붙은 경우도 보존.
    if re.search(r"\b20\d{2}\b", text):
        return True
    if re.search(r"\d{1,2}\.\s*\d{1,2}\.", text):
        return True
    return False


def _divider_heading_should_emit_content(line: str, rule_type: str) -> bool:
    if _classify_heading_kind(rule_type) != "divider":
        return False
    remainder = _strip_heading_notation(line, rule_type)
    return bool(remainder and re.search(r"[가-힣A-Za-z0-9]", remainder))


def _dedupe_hwpx_cover_segments(segments: list[ParsedSegment]) -> list[ParsedSegment]:
    if not segments:
        return []

    deduped: list[ParsedSegment] = []
    idx = 0
    while idx < len(segments):
        current = segments[idx]
        current_text = _normalize_text(current.content)
        next_segment = segments[idx + 1] if idx + 1 < len(segments) else None
        if (
            next_segment is not None
            and current.block_type == "paragraph"
            and next_segment.block_type == "table"
            and _normalize_text(next_segment.content) == current_text
            and current_text
        ):
            deduped.append(next_segment)
            idx += 2
            continue
        deduped.append(current)
        idx += 1
    return deduped


def _serialize_heading_nodes(stack: list[dict[str, Any]]) -> list[HeadingNode]:
    return [
        HeadingNode(
            depth=int(item.get("depth") or 0),
            text=str(item.get("text") or ""),
            kind=str(item.get("kind") or "title"),
            rule_type=str(item.get("rule_type") or ""),
            notation=str(item.get("notation") or ""),
        )
        for item in stack
        if str(item.get("text") or "").strip()
    ]


def _heading_line_for_content(item: dict[str, Any]) -> str:
    text = _normalize_text(str(item.get("text") or ""))
    return text


def _build_pending_heading_lines(pending_stack: list[dict[str, Any]]) -> list[str]:
    return [
        _heading_line_for_content(item)
        for item in pending_stack
        if _heading_line_for_content(item)
    ]


def _next_pending_heading_stack(stack: list[dict[str, Any]]) -> list[dict[str, Any]]:
    if not stack:
        return []
    return [dict(stack[-1])]


def _is_single_depth1_pending_heading(stack: list[dict[str, Any]]) -> bool:
    return (
        len(stack) == 1
        and int(stack[0].get("depth") or 0) == 1
        and bool(_normalize_text(str(stack[0].get("text") or "")))
    )


def _append_standalone_depth1_heading(
    segments: list[ParsedSegment],
    pending_stack: list[dict[str, Any]],
    *,
    section: str,
    sector: str,
    source_order: int,
) -> None:
    if not _is_single_depth1_pending_heading(pending_stack):
        return
    title_text = _normalize_text(str(pending_stack[0].get("text") or ""))
    if not title_text:
        return
    segments.append(
        ParsedSegment(
            source_order=max(0, source_order - 1),
            location="paragraph",
            heading_depth=1,
            heading_text=title_text,
            heading_path=title_text,
            content=title_text,
            block_type="heading-line",
            chunk_loc=_format_chunk_loc(_display_section_label(section), str(source_order), "heading", "1"),
            section=section,
            sector=sector,
            metadata={"sector": sector},
            heading_nodes=_serialize_heading_nodes(pending_stack),
        )
    )


def _merge_pending_heading_content(content: str, pending_stack: list[dict[str, Any]]) -> str:
    lines = _build_pending_heading_lines(pending_stack)
    body = _normalize_text(content)
    if not lines:
        return body
    if body:
        return "\n".join([*lines, body]).strip()
    return "\n".join(lines).strip()


def _strip_duplicate_title_line_from_table(content: str, heading_stack: list[dict[str, Any]]) -> str:
    body_lines = [line.strip() for line in (content or "").splitlines() if line.strip()]
    if len(body_lines) < 2:
        return (content or "").strip()

    title_lines = [
        _normalize_text(str(item.get("text") or ""))
        for item in heading_stack
        if str(item.get("kind") or "title") == "title" and _normalize_text(str(item.get("text") or ""))
    ]
    if not title_lines:
        return "\n".join(body_lines).strip()

    first_line = _normalize_text(body_lines[0])
    last_title = title_lines[-1]
    if first_line == last_title or first_line.startswith(last_title) or last_title.startswith(first_line):
        return "\n".join(body_lines[1:]).strip()
    return "\n".join(body_lines).strip()


def _is_table_segment(segment: ParsedSegment) -> bool:
    return segment.block_type in {"table", "table-row"}


def _is_table_block_type(block_type: str) -> bool:
    return block_type in {"table", "table-row", "table-single-row"}


def _find_table_caption_lines(text: str) -> list[str]:
    if not text:
        return []
    line_pattern = re.compile(r"(?m)^\s*([《<〈「『].+[》>〉」』])\s*$")
    return [match.group(1).strip() for match in line_pattern.finditer(text) if match.group(1).strip()]


def _remove_table_caption_lines(text: str) -> str:
    if not text:
        return ""
    cleaned_lines = [
        line
        for line in (text or "").splitlines()
        if not _find_table_caption_lines(line)
    ]
    cleaned = "\n".join(cleaned_lines).strip()
    cleaned = re.sub(r"\n{3,}", "\n\n", cleaned).strip()
    return cleaned


def _find_previous_table_segment_index(segments: list[ParsedSegment], start_idx: int) -> int | None:
    current_heading = str(segments[start_idx].heading_path or "")
    for idx in range(start_idx - 1, -1, -1):
        candidate = segments[idx]
        if current_heading and candidate.heading_path != current_heading:
            break
        if _is_table_segment(candidate):
            return idx
        if candidate.block_type not in {"paragraph", "heading-inline"}:
            break
    return None


def _find_next_table_segment_index(segments: list[ParsedSegment], start_idx: int) -> int | None:
    current_heading = str(segments[start_idx].heading_path or "")
    for idx in range(start_idx + 1, len(segments)):
        candidate = segments[idx]
        if current_heading and candidate.heading_path != current_heading:
            break
        if _is_table_segment(candidate):
            return idx
        if candidate.block_type not in {"paragraph", "heading-inline"}:
            break
    return None


def _find_table_cluster_start(segments: list[ParsedSegment], table_idx: int) -> int:
    current_heading = str(segments[table_idx].heading_path or "")
    idx = table_idx
    while idx > 0:
        prev = segments[idx - 1]
        if not _is_table_segment(prev):
            break
        if current_heading and prev.heading_path != current_heading:
            break
        idx -= 1
    return idx


def _apply_table_caption(segment: ParsedSegment, caption: str, *, position: str) -> None:
    normalized_caption = _normalize_text(caption)
    if not normalized_caption:
        return
    metadata = {**(segment.metadata or {})}
    existing = _normalize_text(str(metadata.get("table_caption") or ""))
    if existing and existing != normalized_caption:
        normalized_caption = f"{existing}\n{normalized_caption}"
    metadata["table_caption"] = normalized_caption
    metadata["table_caption_position"] = position
    segment.metadata = metadata
    if normalized_caption not in (segment.content or ""):
        segment.content = f"{normalized_caption}\n{segment.content}".strip()


def _resolve_table_caption_target(
    segments: list[ParsedSegment],
    caption_segment_idx: int,
) -> tuple[int | None, str | None]:
    prev_idx = _find_previous_table_segment_index(segments, caption_segment_idx)
    next_idx = _find_next_table_segment_index(segments, caption_segment_idx)

    if next_idx is not None:
        return _find_table_cluster_start(segments, next_idx), "before-table"
    if prev_idx is not None:
        return _find_table_cluster_start(segments, prev_idx), "after-table"
    return None, None


def _attach_table_captions(segments: list[ParsedSegment]) -> list[ParsedSegment]:
    if not segments:
        return segments

    for idx, segment in enumerate(segments):
        if segment.block_type not in {"paragraph", "heading-inline", "table", "table-row", "table-single-row"}:
            continue
        if bool((segment.metadata or {}).get("table_caption_embedded")):
            continue

        captions = _find_table_caption_lines(segment.content)
        if not captions:
            continue

        target_idx, position = _resolve_table_caption_target(segments, idx)
        if target_idx is None or not position:
            continue

        segment.content = _remove_table_caption_lines(segment.content)
        for caption in captions:
            _apply_table_caption(segments[target_idx], caption, position=position)

    return [segment for segment in segments if (segment.content or "").strip()]


def _format_chunk_loc(*parts: str) -> str:
    cleaned = [str(part).strip() for part in parts if str(part).strip()]
    return ":".join(cleaned)


def _display_section_label(section: str) -> str:
    raw = str(section or "").strip()
    if not raw:
        return "section"
    match = re.search(r"section(\d+)\.xml$", raw, flags=re.IGNORECASE)
    if match:
        return f"s{match.group(1)}"
    return raw


def parse_markdown_document(*, file_id: str, filename: str, content: bytes, heading_schema: dict[str, Any]) -> list[ParsedSegment]:
    try:
        text = content.decode("utf-8")
    except UnicodeDecodeError:
        text = content.decode("utf-8", errors="ignore")

    compiled_levels = _compile_heading_levels(heading_schema)
    heading_stack: list[dict[str, Any]] = []
    pending_heading_stack: list[dict[str, Any]] = []
    segments: list[ParsedSegment] = []
    source_order = 0

    for raw_line in text.splitlines():
        line = _normalize_text(raw_line)
        if not line:
            continue
        md_heading = _extract_markdown_heading(line)
        heading_depth: int | None = None
        heading_text = ""
        matched_level = None
        if md_heading:
            heading_depth, heading_text = md_heading
        else:
            matched_level = _match_heading_level(line, "paragraph", compiled_levels)
            if matched_level is not None:
                heading_depth = int(matched_level["depth"])
                heading_text = line

        if heading_depth is not None:
            matched_rule_type = str((matched_level or {}).get("rule_type") or "")
            heading_stack = _update_heading_stack(heading_stack, heading_depth, heading_text, rule_type=matched_rule_type)
            pending_heading_stack = [dict(item) for item in heading_stack]
            if _heading_has_inline_payload(line) or _divider_heading_should_emit_content(line, matched_rule_type):
                source_order += 1
                heading_path = _build_heading_path(heading_stack)
                segments.append(
                    ParsedSegment(
                        source_order=source_order,
                        location="paragraph",
                        heading_depth=heading_stack[-1]["depth"] if heading_stack else None,
                        heading_text=heading_stack[-1]["text"] if heading_stack else "",
                        heading_path=heading_path,
                        content=_merge_pending_heading_content(line, []),
                        block_type="heading-inline",
                        chunk_loc=_format_chunk_loc("line", str(source_order)),
                        section=filename,
                        sector="main",
                        heading_nodes=_serialize_heading_nodes(heading_stack),
                    )
                )
                pending_heading_stack = []
            continue

        source_order += 1
        heading_path = _build_heading_path(heading_stack)
        chunk_loc = _format_chunk_loc("line", str(source_order))
        segments.append(
            ParsedSegment(
                source_order=source_order,
                location="paragraph",
                heading_depth=heading_stack[-1]["depth"] if heading_stack else None,
                heading_text=heading_stack[-1]["text"] if heading_stack else "",
                heading_path=heading_path,
                content=_merge_pending_heading_content(line, pending_heading_stack),
                block_type="paragraph",
                chunk_loc=chunk_loc,
                section=filename,
                sector="main",
                heading_nodes=_serialize_heading_nodes(heading_stack),
            )
        )
        pending_heading_stack = []
    return segments


def parse_docx_document(*, file_id: str, filename: str, content: bytes, heading_schema: dict[str, Any]) -> list[ParsedSegment]:
    try:
        from docx import Document  # type: ignore[import-not-found]
    except ImportError as exc:
        raise RagEmbeddingParseError("python-docx 모듈이 설치되지 않았습니다.") from exc

    try:
        document = Document(io.BytesIO(content))
    except Exception as exc:  # noqa: BLE001
        raise RagEmbeddingParseError(f"DOCX 를 열 수 없습니다: {exc}") from exc

    compiled_levels = _compile_heading_levels(heading_schema)
    heading_stack: list[dict[str, Any]] = []
    pending_heading_stack: list[dict[str, Any]] = []
    segments: list[ParsedSegment] = []
    source_order = 0

    for para in document.paragraphs:
        line = _normalize_text(para.text or "")
        if not line:
            continue
        style_name = (para.style.name if para.style else "") or ""
        heading_matches = re.findall(r"\d+", style_name)
        heading_depth = int(heading_matches[0]) if style_name.lower().startswith("heading") and heading_matches else None
        matched_level = None
        if heading_depth is None:
            matched_level = _match_heading_level(line, "paragraph", compiled_levels)
            if matched_level is not None:
                heading_depth = int(matched_level["depth"])
        if heading_depth is not None:
            matched_rule_type = str((matched_level or {}).get("rule_type") or "")
            heading_stack = _update_heading_stack(heading_stack, heading_depth, line, rule_type=matched_rule_type)
            pending_heading_stack = [dict(item) for item in heading_stack]
            if _heading_has_inline_payload(line) or _divider_heading_should_emit_content(line, matched_rule_type):
                source_order += 1
                heading_path = _build_heading_path(heading_stack)
                segments.append(
                    ParsedSegment(
                        source_order=source_order,
                        location="paragraph",
                        heading_depth=heading_stack[-1]["depth"] if heading_stack else None,
                        heading_text=heading_stack[-1]["text"] if heading_stack else "",
                        heading_path=heading_path,
                        content=_merge_pending_heading_content(line, []),
                        block_type="heading-inline",
                        chunk_loc=_format_chunk_loc("paragraph", str(source_order)),
                        section="body",
                        sector="main",
                        heading_nodes=_serialize_heading_nodes(heading_stack),
                    )
                )
                pending_heading_stack = []
            continue
        source_order += 1
        heading_path = _build_heading_path(heading_stack)
        segments.append(
            ParsedSegment(
                source_order=source_order,
                location="paragraph",
                heading_depth=heading_stack[-1]["depth"] if heading_stack else None,
                heading_text=heading_stack[-1]["text"] if heading_stack else "",
                heading_path=heading_path,
                content=_merge_pending_heading_content(line, pending_heading_stack),
                block_type="paragraph",
                chunk_loc=_format_chunk_loc("paragraph", str(source_order)),
                section="body",
                sector="main",
                heading_nodes=_serialize_heading_nodes(heading_stack),
            )
        )
        pending_heading_stack = []

    for table_index, table in enumerate(document.tables, start=1):
        matrix: list[list[str]] = []
        for row in table.rows:
            cells = [_normalize_text(cell.text or "") for cell in row.cells]
            if any(cells):
                matrix.append(cells)
        rows = [" | ".join(cell for cell in row if cell) for row in matrix]
        if not rows:
            continue
        if len(matrix) == 1:
            flattened = _flatten_table_row(matrix[0])
            if flattened:
                matched_level = _match_heading_level(flattened, "table", compiled_levels)
                heading_depth = int(matched_level["depth"]) if matched_level is not None else None
                if heading_depth is not None:
                    matched_rule_type = str((matched_level or {}).get("rule_type") or "")
                    heading_stack = _update_heading_stack(heading_stack, heading_depth, flattened, rule_type=matched_rule_type)
                    pending_heading_stack = [dict(item) for item in heading_stack]
                    if _heading_has_inline_payload(flattened) or _divider_heading_should_emit_content(flattened, matched_rule_type):
                        source_order += 1
                        heading_path = _build_heading_path(heading_stack)
                        segments.append(
                            ParsedSegment(
                                source_order=source_order,
                                location="table",
                                heading_depth=heading_stack[-1]["depth"] if heading_stack else None,
                                heading_text=heading_stack[-1]["text"] if heading_stack else "",
                                heading_path=heading_path,
                                content=_merge_pending_heading_content(flattened, []),
                                block_type="heading-inline",
                                chunk_loc=_format_chunk_loc("table", str(table_index)),
                                section="body",
                                sector="main",
                                heading_nodes=_serialize_heading_nodes(heading_stack),
                            )
                        )
                        pending_heading_stack = []
                    continue
        source_order += 1
        heading_path = _build_heading_path(heading_stack)
        segments.append(
            ParsedSegment(
                source_order=source_order,
                location="table",
                heading_depth=heading_stack[-1]["depth"] if heading_stack else None,
                heading_text=heading_stack[-1]["text"] if heading_stack else "",
                heading_path=heading_path,
                content="\n".join(rows),
                block_type="table",
                chunk_loc=_format_chunk_loc("table", str(table_index)),
                section="table",
                sector="main",
                metadata={"row_count": len(rows)},
                heading_nodes=_serialize_heading_nodes(heading_stack),
            )
        )
        pending_heading_stack = []
    return segments


def parse_pdf_document(*, file_id: str, filename: str, content: bytes, heading_schema: dict[str, Any]) -> list[ParsedSegment]:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise RagEmbeddingParseError("pypdf 모듈이 설치되지 않았습니다.") from exc

    try:
        reader = PdfReader(io.BytesIO(content))
    except Exception as exc:  # noqa: BLE001
        raise RagEmbeddingParseError(f"PDF 를 열 수 없습니다: {exc}") from exc

    compiled_levels = _compile_heading_levels(heading_schema)
    heading_stack: list[dict[str, Any]] = []
    pending_heading_stack: list[dict[str, Any]] = []
    segments: list[ParsedSegment] = []
    source_order = 0

    for page_index, page in enumerate(reader.pages, start=1):
        try:
            raw_text = page.extract_text() or ""
        except Exception:
            continue
        for para in re.split(r"\n{2,}", raw_text):
            line = _normalize_text(para)
            if not line:
                continue
            matched_level = _match_heading_level(line, "paragraph", compiled_levels)
            heading_depth = int(matched_level["depth"]) if matched_level is not None else None
            if heading_depth is not None:
                matched_rule_type = str((matched_level or {}).get("rule_type") or "")
                heading_stack = _update_heading_stack(heading_stack, heading_depth, line, rule_type=matched_rule_type)
                pending_heading_stack = [dict(item) for item in heading_stack]
                if _heading_has_inline_payload(line) or _divider_heading_should_emit_content(line, matched_rule_type):
                    source_order += 1
                    heading_path = _build_heading_path(heading_stack)
                    segments.append(
                        ParsedSegment(
                            source_order=source_order,
                            location="paragraph",
                            heading_depth=heading_stack[-1]["depth"] if heading_stack else None,
                            heading_text=heading_stack[-1]["text"] if heading_stack else "",
                            heading_path=heading_path,
                            content=_merge_pending_heading_content(line, []),
                            block_type="heading-inline",
                            chunk_loc=_format_chunk_loc("page", str(page_index), "para", str(source_order)),
                            section=f"page_{page_index}",
                            sector="main",
                            heading_nodes=_serialize_heading_nodes(heading_stack),
                        )
                    )
                    pending_heading_stack = []
                continue
            source_order += 1
            heading_path = _build_heading_path(heading_stack)
            segments.append(
                ParsedSegment(
                    source_order=source_order,
                    location="paragraph",
                    heading_depth=heading_stack[-1]["depth"] if heading_stack else None,
                    heading_text=heading_stack[-1]["text"] if heading_stack else "",
                    heading_path=heading_path,
                    content=_merge_pending_heading_content(line, pending_heading_stack),
                    block_type="paragraph",
                    chunk_loc=_format_chunk_loc("page", str(page_index), "para", str(source_order)),
                    section=f"page_{page_index}",
                    sector="main",
                    heading_nodes=_serialize_heading_nodes(heading_stack),
                )
            )
            pending_heading_stack = []
    return segments


def parse_hwpx_document(
    *,
    file_id: str,
    filename: str,
    content: bytes,
    heading_schema: dict[str, Any],
    appendix_schema: dict[str, Any] | None = None,
    body_exit_criteria: dict[str, Any] | None = None,
    appendix_exit_criteria: dict[str, Any] | None = None,
) -> list[ParsedSegment]:
    compiled_levels = _compile_heading_levels(heading_schema)
    compiled_appendix_levels = _compile_heading_levels(appendix_schema or {}) if isinstance(appendix_schema, dict) else []
    events = _extract_hwpx_events(content)
    heading_stacks: dict[str, list[dict[str, Any]]] = {"main": [], "appendix": []}
    pending_heading_stacks: dict[str, list[dict[str, Any]]] = {"main": [], "appendix": []}
    segments: list[ParsedSegment] = []
    mode = "main"

    for event in events:
        text = _normalize_text(str(event.get("text") or ""))
        match_text = _normalize_text(str(event.get("match_text") or text))
        if not text or not match_text:
            continue
        source_order = int(event["source_order"])
        section = str(event.get("section") or "")
        location = _normalize_event_location(event.get("location") or "paragraph")
        event_metadata = event.get("metadata") if isinstance(event.get("metadata"), dict) else None

        if compiled_appendix_levels:
            matched_main_level = _match_heading_level(match_text, location, compiled_levels, metadata=event_metadata)
            matched_appendix_level = _match_heading_level(match_text, location, compiled_appendix_levels, metadata=event_metadata)

            if mode == "main":
                matched_body_exit = _matches_exit_criteria_runtime(
                    match_text,
                    location,
                    criteria=body_exit_criteria,
                    compiled_levels=compiled_appendix_levels,
                    metadata=event_metadata,
                )
                if matched_body_exit or (
                    matched_appendix_level is not None and int(matched_appendix_level.get("depth") or 0) == 1
                ):
                    mode = "appendix"

            else:
                matched_main_reentry = _is_main_reentry_level(matched_main_level)
                matched_appendix_exit = (
                    _matches_exit_criteria_runtime(
                        match_text,
                        location,
                        criteria=appendix_exit_criteria,
                        compiled_levels=compiled_levels,
                        metadata=event_metadata,
                        exclude_symbolic_reentry_levels=True,
                    )
                    if isinstance(appendix_exit_criteria, dict)
                    else _matches_default_appendix_exit_runtime(match_text, location, compiled_levels, metadata=event_metadata)
                )
                if matched_main_reentry or matched_appendix_exit:
                    mode = "main"

        active_compiled_levels = compiled_appendix_levels if mode == "appendix" else compiled_levels
        active_heading_stack = heading_stacks[mode]
        active_pending_heading_stack = pending_heading_stacks[mode]
        matched_level = _match_heading_level(match_text, location, active_compiled_levels, metadata=event_metadata)
        heading_depth = int(matched_level["depth"]) if matched_level is not None else None
        if heading_depth is not None:
            previous_pending_stack = [dict(item) for item in active_pending_heading_stack]
            if _is_single_depth1_pending_heading(previous_pending_stack) and heading_depth > 1:
                _append_standalone_depth1_heading(
                    segments,
                    previous_pending_stack,
                    section=section,
                    sector=mode,
                    source_order=source_order,
                )
            matched_rule_type = str((matched_level or {}).get("rule_type") or "")
            active_heading_stack = _update_heading_stack(active_heading_stack, heading_depth, match_text, rule_type=matched_rule_type)
            if mode == "appendix" and heading_depth == 1 and active_heading_stack:
                active_heading_stack[-1]["kind"] = "title"
            active_pending_heading_stack = _next_pending_heading_stack(active_heading_stack)
            heading_stacks[mode] = active_heading_stack
            pending_heading_stacks[mode] = active_pending_heading_stack
            if _heading_has_inline_payload(match_text) or _divider_heading_should_emit_content(match_text, matched_rule_type):
                inline_content = _merge_pending_heading_content(html.unescape(text), [])
                segments.append(
                    ParsedSegment(
                        source_order=source_order,
                        location=location,
                        heading_depth=active_heading_stack[-1]["depth"] if active_heading_stack else None,
                        heading_text=active_heading_stack[-1]["text"] if active_heading_stack else "",
                        heading_path=_build_heading_path_for_content(active_heading_stack, inline_content),
                        content=inline_content,
                        block_type="heading-inline",
                        chunk_loc=_format_chunk_loc(_display_section_label(section), str(source_order)),
                        section=section,
                        sector=mode,
                        metadata={"sector": mode},
                        heading_nodes=_serialize_heading_nodes(active_heading_stack),
                    )
                )
                pending_heading_stacks[mode] = []
            continue

        block_type = str(event.get("block_type") or "paragraph")
        chunk_loc = _format_chunk_loc(_display_section_label(section), str(source_order))
        metadata = event.get("metadata")
        merged_metadata = {
            **(metadata if isinstance(metadata, dict) else {}),
            "sector": mode,
        }
        content_text = html.unescape(text)
        if _is_single_depth1_pending_heading(active_pending_heading_stack):
            _append_standalone_depth1_heading(
                segments,
                active_pending_heading_stack,
                section=section,
                sector=mode,
                source_order=source_order,
            )
            active_pending_heading_stack = []
            pending_heading_stacks[mode] = []
        if _is_table_block_type(block_type):
            content_text = _merge_pending_heading_content(content_text, [])
            if block_type == "table-single-row":
                content_text = _strip_duplicate_title_line_from_table(content_text, active_heading_stack)
        else:
            content_text = _merge_pending_heading_content(content_text, active_pending_heading_stack)
        heading_path = _build_heading_path_for_content(active_heading_stack, content_text)
        segments.append(
            ParsedSegment(
                source_order=source_order,
                location=location,
                heading_depth=active_heading_stack[-1]["depth"] if active_heading_stack else None,
                heading_text=active_heading_stack[-1]["text"] if active_heading_stack else "",
                heading_path=heading_path,
                content=content_text,
                block_type=block_type,
                chunk_loc=chunk_loc,
                section=section,
                sector=mode,
                metadata=merged_metadata,
                heading_nodes=_serialize_heading_nodes(active_heading_stack),
            )
        )
        pending_heading_stacks[mode] = []
    return _attach_table_captions(_dedupe_hwpx_cover_segments(segments))


def parse_document(
    *,
    file_id: str,
    filename: str,
    content: bytes,
    heading_schema: dict[str, Any],
    appendix_schema: dict[str, Any] | None = None,
    body_exit_criteria: dict[str, Any] | None = None,
    appendix_exit_criteria: dict[str, Any] | None = None,
) -> list[ParsedSegment]:
    suffix = Path(filename).suffix.lower()
    if suffix == ".md":
        return parse_markdown_document(file_id=file_id, filename=filename, content=content, heading_schema=heading_schema)
    if suffix == ".docx":
        return parse_docx_document(file_id=file_id, filename=filename, content=content, heading_schema=heading_schema)
    if suffix == ".pdf":
        return parse_pdf_document(file_id=file_id, filename=filename, content=content, heading_schema=heading_schema)
    if suffix == ".hwpx":
        return parse_hwpx_document(
            file_id=file_id,
            filename=filename,
            content=content,
            heading_schema=heading_schema,
            appendix_schema=appendix_schema,
            body_exit_criteria=body_exit_criteria,
            appendix_exit_criteria=appendix_exit_criteria,
        )
    raise RagEmbeddingParseError(f"지원하지 않는 형식입니다: {suffix}")


class HwpxParser:
    """RAG embedding 용 문서 parser 진입점."""

    def parse_document(
        self,
        *,
        file_id: str,
        filename: str,
        content: bytes,
        heading_schema: dict[str, Any],
        appendix_schema: dict[str, Any] | None = None,
        body_exit_criteria: dict[str, Any] | None = None,
        appendix_exit_criteria: dict[str, Any] | None = None,
    ) -> list[ParsedSegment]:
        """파일 형식에 맞는 parser 를 호출해 ParsedSegment 목록을 만든다."""
        return parse_document(
            file_id=file_id,
            filename=filename,
            content=content,
            heading_schema=heading_schema,
            appendix_schema=appendix_schema,
            body_exit_criteria=body_exit_criteria,
            appendix_exit_criteria=appendix_exit_criteria,
        )


__all__ = [
    "HeadingNode",
    "ParsedSegment",
    "RagEmbeddingParseError",
    "HwpxParser",
    "parse_document",
]
