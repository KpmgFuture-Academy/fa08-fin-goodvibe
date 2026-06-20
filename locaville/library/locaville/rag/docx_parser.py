from __future__ import annotations

import io
import re
from typing import Any

from .document_models import ParsedSegment, RagEmbeddingParseError
from .hwpx_parser import (
    _build_heading_path,
    _compile_heading_levels,
    _divider_heading_should_emit_content,
    _flatten_table_row,
    _format_chunk_loc,
    _heading_has_inline_payload,
    _match_heading_level,
    _merge_pending_heading_content,
    _normalize_text,
    _serialize_heading_nodes,
    _update_heading_stack,
)


def parse_docx_document(
    *,
    file_id: str,
    filename: str,
    content: bytes,
    heading_schema: dict[str, Any],
    appendix_schema: dict[str, Any] | None = None,
    body_exit_criteria: dict[str, Any] | None = None,
    appendix_exit_criteria: dict[str, Any] | None = None,
) -> list[ParsedSegment]:
    del file_id, appendix_schema, body_exit_criteria, appendix_exit_criteria

    try:
        from docx import Document  # type: ignore[import-not-found]
    except ImportError as exc:
        raise RagEmbeddingParseError("python-docx 모듈이 설치되지 않았습니다.") from exc

    try:
        document = Document(io.BytesIO(content))
    except Exception as exc:  # noqa: BLE001
        raise RagEmbeddingParseError(f"DOCX 를 열 수 없습니다: {exc}") from exc

    compiled_levels = _compile_heading_levels(heading_schema)
    heading_stack: list[dict[str, Any]] = []
    pending_heading_stack: list[dict[str, Any]] = []
    segments: list[ParsedSegment] = []
    source_order = 0

    for para in document.paragraphs:
        line = _normalize_text(para.text or "")
        if not line:
            continue
        style_name = (para.style.name if para.style else "") or ""
        heading_matches = re.findall(r"\d+", style_name)
        heading_depth = int(heading_matches[0]) if style_name.lower().startswith("heading") and heading_matches else None
        matched_level = None
        if heading_depth is None:
            matched_level = _match_heading_level(line, "paragraph", compiled_levels)
            if matched_level is not None:
                heading_depth = int(matched_level["depth"])
        if heading_depth is not None:
            matched_rule_type = str((matched_level or {}).get("rule_type") or "")
            heading_stack = _update_heading_stack(heading_stack, heading_depth, line, rule_type=matched_rule_type)
            pending_heading_stack = [dict(item) for item in heading_stack]
            if _heading_has_inline_payload(line) or _divider_heading_should_emit_content(line, matched_rule_type):
                source_order += 1
                heading_path = _build_heading_path(heading_stack)
                segments.append(
                    ParsedSegment(
                        source_order=source_order,
                        location="paragraph",
                        heading_depth=heading_stack[-1]["depth"] if heading_stack else None,
                        heading_text=heading_stack[-1]["text"] if heading_stack else "",
                        heading_path=heading_path,
                        content=_merge_pending_heading_content(line, []),
                        block_type="heading-inline",
                        chunk_loc=_format_chunk_loc("paragraph", str(source_order)),
                        section="body",
                        sector="main",
                        heading_nodes=_serialize_heading_nodes(heading_stack),
                    )
                )
                pending_heading_stack = []
            continue
        source_order += 1
        heading_path = _build_heading_path(heading_stack)
        segments.append(
            ParsedSegment(
                source_order=source_order,
                location="paragraph",
                heading_depth=heading_stack[-1]["depth"] if heading_stack else None,
                heading_text=heading_stack[-1]["text"] if heading_stack else "",
                heading_path=heading_path,
                content=_merge_pending_heading_content(line, pending_heading_stack),
                block_type="paragraph",
                chunk_loc=_format_chunk_loc("paragraph", str(source_order)),
                section="body",
                sector="main",
                heading_nodes=_serialize_heading_nodes(heading_stack),
            )
        )
        pending_heading_stack = []

    for table_index, table in enumerate(document.tables, start=1):
        matrix: list[list[str]] = []
        for row in table.rows:
            cells = [_normalize_text(cell.text or "") for cell in row.cells]
            if any(cells):
                matrix.append(cells)
        rows = [" | ".join(cell for cell in row if cell) for row in matrix]
        if not rows:
            continue
        if len(matrix) == 1:
            flattened = _flatten_table_row(matrix[0])
            if flattened:
                matched_level = _match_heading_level(flattened, "table", compiled_levels)
                heading_depth = int(matched_level["depth"]) if matched_level is not None else None
                if heading_depth is not None:
                    matched_rule_type = str((matched_level or {}).get("rule_type") or "")
                    heading_stack = _update_heading_stack(heading_stack, heading_depth, flattened, rule_type=matched_rule_type)
                    pending_heading_stack = [dict(item) for item in heading_stack]
                    if _heading_has_inline_payload(flattened) or _divider_heading_should_emit_content(flattened, matched_rule_type):
                        source_order += 1
                        heading_path = _build_heading_path(heading_stack)
                        segments.append(
                            ParsedSegment(
                                source_order=source_order,
                                location="table",
                                heading_depth=heading_stack[-1]["depth"] if heading_stack else None,
                                heading_text=heading_stack[-1]["text"] if heading_stack else "",
                                heading_path=heading_path,
                                content=_merge_pending_heading_content(flattened, []),
                                block_type="heading-inline",
                                chunk_loc=_format_chunk_loc("table", str(table_index)),
                                section="body",
                                sector="main",
                                heading_nodes=_serialize_heading_nodes(heading_stack),
                            )
                        )
                        pending_heading_stack = []
                    continue
        source_order += 1
        heading_path = _build_heading_path(heading_stack)
        segments.append(
            ParsedSegment(
                source_order=source_order,
                location="table",
                heading_depth=heading_stack[-1]["depth"] if heading_stack else None,
                heading_text=heading_stack[-1]["text"] if heading_stack else "",
                heading_path=heading_path,
                content="\n".join(rows),
                block_type="table",
                chunk_loc=_format_chunk_loc("table", str(table_index)),
                section="table",
                sector="main",
                metadata={"row_count": len(rows)},
                heading_nodes=_serialize_heading_nodes(heading_stack),
            )
        )
        pending_heading_stack = []
    return segments


class DocxDocumentParser:
    def parse_document(
        self,
        *,
        file_id: str,
        filename: str,
        content: bytes,
        heading_schema: dict[str, Any],
        appendix_schema: dict[str, Any] | None = None,
        body_exit_criteria: dict[str, Any] | None = None,
        appendix_exit_criteria: dict[str, Any] | None = None,
    ) -> list[ParsedSegment]:
        return parse_docx_document(
            file_id=file_id,
            filename=filename,
            content=content,
            heading_schema=heading_schema,
            appendix_schema=appendix_schema,
            body_exit_criteria=body_exit_criteria,
            appendix_exit_criteria=appendix_exit_criteria,
        )


__all__ = [
    "DocxDocumentParser",
    "parse_docx_document",
]
