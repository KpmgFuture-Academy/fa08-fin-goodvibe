"""정책 문서 (PDF / DOCX / HWPX) → 청크 → Supabase pgvector 영구 인덱싱.

확장자 별로 텍스트 추출기를 분기 (PDF=pypdf, DOCX=python-docx, HWPX=기존 hwpx_ingest).
공통 ParsedBlock list 로 정규화 → hwpx_ingest._chunk_parsed_blocks 재사용해 청크 →
supabase_rag_service.{embed_texts, chunk_to_row, insert_rows} 로 rag_chunks 테이블에 영구 저장.

사용처:
  - 사업 등록 화면에서 시행령 업로드 → 자동 청킹 + 임베딩 + 사업 초안 생성

NOTE: 파일 자체는 임시 메모리(bytes)에서만 다루며 디스크에 영구 저장하지 않는다.
      재학습이 필요하면 사용자가 같은 파일을 다시 업로드. RAG 청크만 DB 에 영구 보관.
"""
from __future__ import annotations

import io
import re
from dataclasses import asdict
from pathlib import Path
from typing import Any

from app.services.hwpx_ingest_service import (
    ParsedBlock,
    _chunk_parsed_blocks,
    _read_hwpx_paragraph_blocks,
)
from app.services.supabase_rag_service import (
    chunk_to_row,
    embed_texts,
    improve_chunks,
    insert_rows,
)


class DocumentIngestError(Exception):
    """문서 ingest 실패. 라우터가 HTTP 400 으로 변환."""


MAX_DOCUMENT_BYTES = 30 * 1024 * 1024  # 30MB
SUPPORTED_SUFFIXES = (".pdf", ".docx", ".hwpx")


def _normalize_paragraph(text: str) -> str:
    """줄바꿈/탭 압축. 빈 줄 제거."""
    if not text:
        return ""
    lines = [" ".join(line.split()) for line in text.splitlines()]
    return "\n".join(line for line in lines if line.strip()).strip()


def _is_heading(text: str) -> bool:
    """공통 heading 패턴 — hwpx_ingest 의 패턴 + PDF/DOCX 의 자주 등장 패턴."""
    t = (text or "").strip()
    if not t or len(t) > 120:
        return False
    patterns = (
        r"^[ⅠⅡⅢⅣⅤⅥⅦⅧⅨⅩ]\s+.+",
        r"^[IVX]+\.\s+.+",
        r"^제\s*\d+\s*[장절관조]\b",
        r"^\d+\.\s+\S",
        r"^\d+\)\s+\S",
        r"^[가-힣]\.\s+",
        r"^[가나다라마바사아자차카타파하]\.\s+",
        r"^[①②③④⑤⑥⑦⑧⑨⑩⑪⑫⑬⑭⑮⑯⑰⑱⑲⑳]\s*",
        r"^[□■◇◆○●▣▶◀]\s*",
    )
    return any(re.match(p, t) for p in patterns)


# ============================================================
# PDF
# ============================================================

def _extract_pdf_blocks(content: bytes, filename: str) -> list[ParsedBlock]:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise DocumentIngestError(
            "pypdf 모듈이 설치되지 않았습니다. requirements.txt 의 pypdf 를 설치해 주세요."
        ) from exc

    try:
        reader = PdfReader(io.BytesIO(content))
    except Exception as exc:  # noqa: BLE001
        raise DocumentIngestError(f"PDF 를 열 수 없습니다: {exc}") from exc

    blocks: list[ParsedBlock] = []
    current_heading = ""

    for page_idx, page in enumerate(reader.pages, start=1):
        try:
            raw = page.extract_text() or ""
        except Exception:  # noqa: BLE001
            continue
        # PDF 는 줄바꿈이 들쭉날쭉 — 빈 줄 두 개로 단락 구분 우선.
        paragraphs = re.split(r"\n{2,}", raw)
        for para in paragraphs:
            text = _normalize_paragraph(para)
            if not text:
                continue
            if _is_heading(text):
                current_heading = text
                continue
            blocks.append(
                ParsedBlock(
                    title=current_heading,
                    text=text,
                    section=f"page_{page_idx}",
                )
            )
    if not blocks:
        raise DocumentIngestError(f"{filename} 에서 파싱 가능한 텍스트를 찾지 못했어요.")
    return blocks


# ============================================================
# DOCX
# ============================================================

def _extract_docx_blocks(content: bytes, filename: str) -> list[ParsedBlock]:
    try:
        from docx import Document  # type: ignore[import-not-found]
    except ImportError as exc:
        raise DocumentIngestError(
            "python-docx 모듈이 설치되지 않았습니다. requirements.txt 의 python-docx 를 설치해 주세요."
        ) from exc

    try:
        doc = Document(io.BytesIO(content))
    except Exception as exc:  # noqa: BLE001
        raise DocumentIngestError(f"DOCX 를 열 수 없습니다: {exc}") from exc

    blocks: list[ParsedBlock] = []
    current_heading = ""

    for para in doc.paragraphs:
        text = _normalize_paragraph(para.text or "")
        if not text:
            continue
        style_name = (para.style.name if para.style else "") or ""
        # DOCX 의 Heading 1/2/3 style 또는 굵은 문장은 heading 으로 취급.
        if style_name.lower().startswith("heading") or _is_heading(text):
            current_heading = text
            continue
        blocks.append(ParsedBlock(title=current_heading, text=text, section=""))

    # 표 셀의 텍스트도 본문에 합쳐 청킹 (시행령 표 자주 등장).
    for tbl in doc.tables:
        for row in tbl.rows:
            row_text_parts = [
                _normalize_paragraph(cell.text or "") for cell in row.cells
            ]
            row_text = " | ".join(part for part in row_text_parts if part)
            if not row_text:
                continue
            blocks.append(ParsedBlock(title=current_heading, text=row_text, section="table"))

    if not blocks:
        raise DocumentIngestError(f"{filename} 에서 파싱 가능한 텍스트를 찾지 못했어요.")
    return blocks


# ============================================================
# HWPX — 기존 모듈 재사용 (메모리 bytes 지원 위해 임시 파일 거침)
# ============================================================

def _extract_hwpx_blocks(content: bytes, filename: str) -> list[ParsedBlock]:
    """hwpx_ingest._read_hwpx_paragraph_blocks 는 Path 를 받음 → 임시파일 경유."""
    import tempfile

    with tempfile.NamedTemporaryFile(suffix=".hwpx", delete=False) as tmp:
        tmp.write(content)
        tmp_path = Path(tmp.name)
    try:
        blocks = _read_hwpx_paragraph_blocks(tmp_path)
    except Exception as exc:  # noqa: BLE001
        raise DocumentIngestError(f"HWPX 를 파싱할 수 없습니다: {exc}") from exc
    finally:
        try:
            tmp_path.unlink(missing_ok=True)
        except Exception:  # noqa: BLE001
            pass

    if not blocks:
        raise DocumentIngestError(f"{filename} 에서 파싱 가능한 텍스트를 찾지 못했어요.")
    return blocks


# ============================================================
# Public API — extract → chunk → embed → INSERT
# ============================================================

def extract_blocks(filename: str, content: bytes) -> list[ParsedBlock]:
    """파일 확장자에 따라 적절한 파서로 ParsedBlock list 생성."""
    if len(content) > MAX_DOCUMENT_BYTES:
        raise DocumentIngestError(
            f"파일이 너무 큽니다 ({len(content) // 1024 // 1024}MB). "
            f"30MB 이하만 업로드해 주세요."
        )
    suffix = Path(filename).suffix.lower()
    if suffix == ".pdf":
        return _extract_pdf_blocks(content, filename)
    if suffix == ".docx":
        return _extract_docx_blocks(content, filename)
    if suffix == ".hwpx":
        return _extract_hwpx_blocks(content, filename)
    raise DocumentIngestError(
        f"지원하지 않는 형식입니다: {suffix}. .pdf / .docx / .hwpx 만 가능합니다."
    )


def ingest_document(
    *,
    filename: str,
    content: bytes,
    chunk_size: int = 800,
) -> dict[str, Any]:
    """파일 bytes → 청크 → 임베딩 → Supabase rag_chunks 영구 적재.

    Returns:
      {
        "filename": str,
        "file_type": "pdf" | "docx" | "hwpx",
        "blocks": int,       # 추출된 단락 수
        "chunks": int,       # 청킹 결과 수 (improve 후)
        "inserted": int,     # DB 에 INSERT/UPDATE 된 행 수
        "chunk_list": list,  # LLM 메타 추출용 청크 (다음 단계에서 사용)
        "preview_blocks": list,  # 초안용 — 앞쪽 단락 일부 (LLM context 용)
      }
    """
    suffix = Path(filename).suffix.lower().lstrip(".")
    blocks = extract_blocks(filename, content)

    raw_chunks = _chunk_parsed_blocks(blocks, source_file=filename, chunk_size=chunk_size)
    chunks = improve_chunks(raw_chunks)

    if not chunks:
        return {
            "filename": filename,
            "file_type": suffix,
            "blocks": len(blocks),
            "chunks": 0,
            "inserted": 0,
            "chunk_list": [],
            "preview_blocks": [asdict(b) for b in blocks[:10]],
        }

    # 임베딩 + INSERT (실패 시 위로 raise — LLM 단계가 RAG 결과를 이용하므로 필수).
    texts = [c.get("text", "") for c in chunks]
    vectors = embed_texts(texts)
    rows = [chunk_to_row(c, v) for c, v in zip(chunks, vectors)]
    inserted = insert_rows(rows)

    return {
        "filename": filename,
        "file_type": suffix,
        "blocks": len(blocks),
        "chunks": len(chunks),
        "inserted": inserted,
        "chunk_list": chunks,
        "preview_blocks": [asdict(b) for b in blocks[:10]],
    }
